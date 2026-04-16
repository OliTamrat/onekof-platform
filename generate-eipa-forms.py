"""
EIPA Software Copyright Forms Generator
=======================================

Generates the four Word documents that accompany the source code deposit
for an Ethiopian Intellectual Property Authority (EIPA) software copyright
filing. The forms are populated with Onekof platform data where known and
leave signature / handwritten fields blank for Oli to complete by hand
before submission.

Output: DEPOSIT/00_FORMS/
    - Form_C-1_Application.docx         (EIPA Form C-1)
    - Software_Description.docx         (Functional description)
    - Publication_Date.docx             (First publication record)
    - Source_Code_Deposit_Cover.docx    (Cover sheet for the 5 source volumes)

Author: Oli Tamrat Oli <oli.oli@udc.edu>
Date:   2026-04-11
"""

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).parent
DEPOSIT_DIR = REPO_ROOT / "DEPOSIT"
FORMS_DIR = DEPOSIT_DIR / "00_FORMS"

TEAL = RGBColor(0x1C, 0x8C, 0x7D)
DARK = RGBColor(0x22, 0x27, 0x2B)
GREY = RGBColor(0x77, 0x77, 0x77)

TODAY = datetime.now().strftime("%B %d, %Y")


# ---------------------------------------------------------------------------
# Document-wide styling helpers
# ---------------------------------------------------------------------------

def set_document_defaults(doc):
    """Set reasonable margins and body font."""
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default style font — Calibri is safe across Word versions; avoid SF Pro.
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def add_title(doc, text, subtitle=None):
    """Big teal title at top of document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = TEAL

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = GREY

    # Teal horizontal rule (empty paragraph with bottom border)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1C8C7D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK


def add_paragraph(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if italic:
        run.italic = True


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
        if not p.runs:
            run = p.add_run(item)
            run.font.size = Pt(11)
        else:
            p.runs[0].text = item


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_field_table(doc, rows, label_width_inches=2.3):
    """
    Two-column table: label | value. Left column shaded teal-tinted.
    `rows` is a list of (label, value) tuples. Use '' for blank fill-in fields
    (we insert a visual underline placeholder).
    """
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.style = 'Light List Accent 1'

    for i, (label, value) in enumerate(rows):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]

        label_cell.width = Inches(label_width_inches)
        value_cell.width = Inches(6.5 - label_width_inches)

        _shade_cell(label_cell, 'E8F5F3')  # light teal tint

        # Label
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lrun = lp.add_run(label)
        lrun.bold = True
        lrun.font.size = Pt(10)
        lrun.font.color.rgb = DARK

        # Value
        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        display = value if value else "__________________________________________"
        vrun = vp.add_run(display)
        vrun.font.size = Pt(10)
        if not value:
            vrun.font.color.rgb = GREY

        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Spacer paragraph after table
    doc.add_paragraph()


def add_checkbox_list(doc, items):
    """Items are list of strings. Uses ☐ character for unchecked box."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run("\u2610  ")  # ballot box
        run.font.size = Pt(12)
        run2 = p.add_run(item)
        run2.font.size = Pt(11)


def add_signature_block(doc):
    add_section_heading(doc, "Declaration & Signature")
    add_paragraph(
        doc,
        "I hereby declare that the information provided in this application is true, "
        "complete, and correct to the best of my knowledge, and that I am the lawful "
        "author and copyright holder of the work described herein.",
    )

    rows = [
        ("Full Name", "Oli Tamrat Oli"),
        ("Title / Capacity", "Founder and Author"),
        ("Email", "oli.oli@udc.edu"),
        ("Signature", ""),
        ("Date", ""),
        ("Place", "Burtonsville, Maryland, USA"),
    ]
    add_field_table(doc, rows)


def add_footer(doc, form_code):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{form_code}  |  Onekof Platform  |  Oli Tamrat Oli  |  Generated {TODAY}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ---------------------------------------------------------------------------
# Form C-1: Software Copyright Application
# ---------------------------------------------------------------------------

def generate_form_c1(output_path):
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "Form C-1")

    add_title(
        doc,
        "FORM C-1",
        "Software Copyright Application — Ethiopian Intellectual Property Authority",
    )

    # --- SECTION 1: APPLICANT ---
    add_section_heading(doc, "Section 1 — Applicant Information")
    add_field_table(doc, [
        ("Full Name of Applicant", "Oli Tamrat Oli"),
        ("Nationality", "Ethiopian"),
        ("Residential Address (per ID)", "605 P St NW, Apt 12"),
        ("City", "Washington"),
        ("State", "District of Columbia (DC)"),
        ("ZIP Code", "20001"),
        ("Country", "United States of America"),
        ("Mailing Address (current)", "14823 Crossvalley Rd"),
        ("Mailing City", "Burtonsville"),
        ("Mailing State", "Maryland (MD)"),
        ("Mailing ZIP Code", "20866"),
        ("Mailing Country", "United States of America"),
        ("Phone", ""),
        ("Email", "oli.oli@udc.edu"),
    ])
    add_paragraph(
        doc,
        "Note: The applicant has recently relocated from Washington, DC to "
        "Burtonsville, Maryland. The residential address above reflects the "
        "address of record on the applicant's US state identification "
        "document, which is pending update. All EIPA correspondence should "
        "be directed to the mailing address listed above.",
        italic=True,
    )

    # --- SECTION 2: AUTHOR ---
    add_section_heading(doc, "Section 2 — Author Information")
    add_field_table(doc, [
        ("Author(s) Full Name", "Oli Tamrat Oli"),
        ("Nationality", "Ethiopian"),
        ("Is author same as applicant?", "Yes"),
        ("Basis of claim (if different)", "N/A — applicant is sole author"),
    ])

    # --- SECTION 3: WORK DETAILS ---
    add_section_heading(doc, "Section 3 — Work Details")
    add_field_table(doc, [
        ("Title of Work", "Onekof Platform"),
        ("Type of Work", "Computer Software (SaaS Platform)"),
        ("Programming Languages",
            "TypeScript, JavaScript, React, SQL, Prisma SDL"),
        ("Framework", "Next.js 14 (App Router)"),
        ("Lines of Source Code", "132,173 lines across 689 files"),
        ("Date of Creation", "February 2026"),
        ("Date of First Publication", "March 2026"),
        ("Country of First Publication", "Ethiopia"),
        ("Medium of Publication", "Web deployment (SaaS) via Vercel"),
    ])

    # --- SECTION 4: DESCRIPTION ---
    add_section_heading(doc, "Section 4 — Description of Work")
    add_paragraph(
        doc,
        "Onekof is an original, multi-tenant enterprise Software-as-a-Service (SaaS) "
        "platform authored by Oli Tamrat Oli for the Ethiopian and East African "
        "market. The platform provides integrated project management, budget "
        "tracking, AI-powered document processing, and native Ethiopian "
        "localization (Ge'ez calendar, Amharic/Oromo/Tigrinya/Somali language "
        "support, and ETB currency with Ethiopian fiscal-year alignment).",
    )
    add_paragraph(
        doc,
        "The software consists of server-side application logic (TypeScript / "
        "Node.js), a React-based client user interface, a PostgreSQL database "
        "schema managed via Prisma ORM, AI document processing algorithms, an "
        "Ethiopian calendar conversion library, and multi-tenant routing and "
        "security middleware. A comprehensive functional description is "
        "provided in the accompanying document \u201cSoftware_Description.docx\u201d.",
    )

    # --- SECTION 5: DEPOSIT MATERIAL ---
    add_section_heading(doc, "Section 5 — Deposit Material")
    add_paragraph(
        doc,
        "The complete source code is submitted on the accompanying CD-ROM. "
        "The deposit contains:",
    )
    add_bullets(doc, [
        "Volume 01 — Pages and API Routes",
        "Volume 02 — React Components",
        "Volume 03 — Libraries, Contexts, Hooks, and Configuration",
        "Volume 04 — Database Schema and Prisma Models",
        "Volume 05 — Configuration, Build Scripts, and Documentation",
        "Full repository archive (onekof-platform-source.zip)",
        "File manifest listing all 689 source files (02_FILE_MANIFEST.docx)",
    ])
    add_paragraph(
        doc,
        "The source code exceeds the 50-page threshold; the full deposit is "
        "provided electronically on the CD-ROM in lieu of the first-25 / "
        "last-25 printed excerpt, per the EIPA electronic submission provision.",
        italic=True,
    )

    # --- SECTION 6: DECLARATION ---
    add_signature_block(doc)

    doc.save(output_path)
    print(f"  OK  {output_path.name}")


# ---------------------------------------------------------------------------
# Software Functionality Description
# ---------------------------------------------------------------------------

def generate_software_description(output_path):
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "Software Description")

    add_title(
        doc,
        "SOFTWARE FUNCTIONALITY DESCRIPTION",
        "Onekof Platform — Companion Document to Form C-1",
    )

    # Title / purpose
    add_section_heading(doc, "Title")
    add_paragraph(
        doc,
        "Onekof Platform \u2014 AI-Powered Enterprise Project Management, Budget "
        "Tracking, and Document Processing Software System.",
    )

    add_section_heading(doc, "Purpose")
    add_paragraph(
        doc,
        "Onekof is a cloud-based, multi-tenant Software-as-a-Service platform "
        "that enables Ethiopian and East African organizations \u2014 including "
        "private companies, government ministries, NGOs, and construction "
        "firms \u2014 to manage projects, budgets, teams, and documents in a "
        "single integrated environment, with native support for Ethiopian "
        "languages, calendar, and fiscal conventions.",
    )

    # Core functionality
    add_section_heading(doc, "Core Functionality")

    add_subheading(doc, "1. Multi-Tenant Architecture")
    add_bullets(doc, [
        "Subdomain-based organization isolation ({orgslug}.onekof.com)",
        "Independent data, users, roles, and configuration per organization",
        "Cross-subdomain authentication via shared JWT cookie scope",
        "Middleware enforces x-organization-slug header for every request",
    ])

    add_subheading(doc, "2. Authentication & Authorization")
    add_bullets(doc, [
        "NextAuth.js v4 with JWT session strategy",
        "Bcrypt password hashing (12 rounds) with email verification",
        "Account lockout protection against brute-force attacks",
        "Role-based access control: Owner, Admin, Manager, Member, Viewer",
        "Project-level visibility: PUBLIC, INTERNAL, PRIVATE, CONFIDENTIAL",
    ])

    add_subheading(doc, "3. Project & Issue Management")
    add_bullets(doc, [
        "Six industry-specific project types: Software, Business, Marketing, "
        "Operations, Research, Construction",
        "Kanban boards, backlog, sprint planning, and issue tracking",
        "Hierarchical tasks, sub-issues, labels, and custom workflows",
        "Optimistic drag-and-drop updates with server rollback on failure",
    ])

    add_subheading(doc, "4. Budget & Financial Management")
    add_bullets(doc, [
        "Multi-level budget allocation and approval workflows",
        "Expense tracking with multi-currency support (ETB default)",
        "Ethiopian fiscal year alignment (Hamle 1 \u2014 Sene 30)",
        "Role-based financial visibility and audit logging",
        "Real-time dashboards for budget utilization and variance",
    ])

    add_subheading(doc, "5. AI-Powered Document Processing")
    add_bullets(doc, [
        "Automated extraction of structured data from invoices, receipts, "
        "contracts, proposals, and government budget documents",
        "Natural language understanding via third-party AI vendor API",
        "Ethiopian-format document support (Amharic text, ETB values, "
        "Ge'ez dates)",
    ])

    add_subheading(doc, "6. Ethiopian-First Localization")
    add_bullets(doc, [
        "Native Ethiopian (Ge'ez) calendar with 13-month support",
        "Bidirectional Gregorian \u2194 Ethiopian date conversion library",
        "Five-language user interface: English, Amharic, Afaan Oromoo, "
        "Tigrinya, Somali",
        "Abyssinica SIL font for Ge'ez script; Inter for Latin scripts",
        "ETB currency as default across all financial screens",
    ])

    add_subheading(doc, "7. Collaboration & Notifications")
    add_bullets(doc, [
        "Team management, member invitations, and role assignment",
        "Activity feeds, comment threads, and @mentions",
        "Watcher and subscription system for issue updates",
        "Document attachments and version tracking",
    ])

    # Technical architecture
    add_section_heading(doc, "Technical Architecture")
    add_field_table(doc, [
        ("Application Framework", "Next.js 14 (App Router)"),
        ("Programming Language", "TypeScript 5.0"),
        ("User Interface", "React 18, Radix UI, Tailwind CSS"),
        ("Database", "PostgreSQL 15 via Prisma ORM 5.0"),
        ("Authentication", "NextAuth.js v4 (JWT)"),
        ("AI Processing", "Third-party AI vendor API"),
        ("Deployment", "Vercel serverless (fra1 region)"),
        ("Monorepo Tooling", "Turborepo + pnpm workspaces"),
        ("Database Hosting", "Supabase (aws-1-eu-central-1)"),
    ], label_width_inches=2.5)

    # Source-code metrics
    add_section_heading(doc, "Source Code Metrics")
    add_field_table(doc, [
        ("Total Source Files", "689"),
        ("Total Lines of Code", "132,173"),
        ("API Routes", "120+ Next.js App Router route handlers"),
        ("React Components", "400+ client and server components"),
        ("Database Models", "80+ Prisma models"),
        ("Locale Files", "5 language JSON dictionaries"),
    ], label_width_inches=2.5)

    # Originality statement
    add_section_heading(doc, "Statement of Originality")
    add_paragraph(
        doc,
        "The Onekof Platform is an original work created by Oli Tamrat Oli. "
        "All source code, database schema, architectural decisions, user "
        "interface designs, business logic, Ethiopian localization assets, "
        "and documentation are original creations. No portion of the work "
        "copies or derives from any other protected work. Third-party "
        "open-source libraries are used under their respective licenses and "
        "are not claimed as part of the work being registered; only the "
        "original application code authored by the applicant is the subject "
        "of this registration.",
    )

    doc.save(output_path)
    print(f"  OK  {output_path.name}")


# ---------------------------------------------------------------------------
# Publication Date
# ---------------------------------------------------------------------------

def generate_publication_date(output_path):
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "Publication Date")

    add_title(
        doc,
        "DATE OF FIRST PUBLICATION",
        "Onekof Platform — Companion Document to Form C-1",
    )

    add_section_heading(doc, "Publication Record")
    add_field_table(doc, [
        ("Title of Work", "Onekof Platform"),
        ("Version at First Publication", "1.0"),
        ("Date of Creation", "February 2026"),
        ("Date of First Publication", "March 2026"),
        ("Method of Publication", "Web deployment (SaaS) via Vercel"),
        ("Country of First Publication", "Ethiopia"),
        ("Primary URL", "https://onekof.com"),
        ("Applicant Email", "oli.oli@udc.edu"),
    ])

    add_section_heading(doc, "Definition of \u201cPublication\u201d for Software")
    add_paragraph(
        doc,
        "Under Ethiopian intellectual property law, \u201cpublication\u201d of a "
        "computer software work refers to the date on which the software was "
        "first made available to the public. For a Software-as-a-Service "
        "product such as Onekof, the applicable publication events include:",
    )
    add_bullets(doc, [
        "First deployment of the application to a publicly accessible "
        "production server",
        "First release of the live application URL to external users",
        "First commercial offering, trial sign-up, or tenant provisioning",
        "First public announcement or demonstration of the platform",
    ])

    add_section_heading(doc, "Evidence Supporting First Publication")
    add_paragraph(
        doc,
        "The following evidence is retained by the applicant and can be "
        "produced upon request by the Ethiopian Intellectual Property "
        "Authority:",
    )
    add_checkbox_list(doc, [
        "Vercel deployment logs showing initial production release",
        "Git commit history from the Onekof platform repository, authored "
        "by Oli Tamrat Oli <oli.oli@udc.edu>",
        "Domain registration records for onekof.com",
        "First tenant organization provisioning records",
        "Supabase database creation timestamps",
        "Source code archive (CD-ROM) submitted with this application",
    ])

    add_section_heading(doc, "Declaration of Publication Date")
    add_paragraph(
        doc,
        "I, Oli Tamrat Oli, declare that the date of first publication stated "
        "above is true and correct, and that the supporting evidence listed "
        "is preserved and available for inspection.",
    )
    add_field_table(doc, [
        ("Declarant", "Oli Tamrat Oli"),
        ("Signature", ""),
        ("Date", ""),
        ("Place", "Burtonsville, Maryland, USA"),
    ])

    doc.save(output_path)
    print(f"  OK  {output_path.name}")


# ---------------------------------------------------------------------------
# Source Code Deposit Cover
# ---------------------------------------------------------------------------

def generate_deposit_cover(output_path):
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "Source Code Deposit Cover")

    add_title(
        doc,
        "SOURCE CODE DEPOSIT",
        "Onekof Platform — CD-ROM Contents and Verification",
    )

    add_section_heading(doc, "Deposit Summary")
    add_field_table(doc, [
        ("Work Title", "Onekof Platform"),
        ("Applicant / Author", "Oli Tamrat Oli"),
        ("Email", "oli.oli@udc.edu"),
        ("Deposit Medium", "CD-ROM (Microsoft Word .docx + ZIP archive)"),
        ("Total Source Files", "689"),
        ("Total Lines of Code", "132,173"),
        ("Deposit Generated", TODAY),
    ])

    add_section_heading(doc, "CD-ROM Contents")
    add_paragraph(
        doc,
        "The accompanying CD-ROM contains the complete source code deposit, "
        "organized as follows:",
    )

    deposit_rows = [
        ("00_FORMS/",
            "EIPA filing forms (Form C-1, Software Description, "
            "Publication Date, this cover sheet)"),
        ("00_COVER_AND_DECLARATION.docx",
            "Authorship declaration and copyright assertion"),
        ("01_PROJECT_DESCRIPTION.docx",
            "Executive summary of the Onekof platform"),
        ("02_FILE_MANIFEST.docx",
            "Complete listing of all 689 deposited source files"),
        ("03_SOURCE_Volume_01_Pages_and_API_Routes.docx",
            "Next.js App Router pages and API route handlers"),
        ("04_SOURCE_Volume_02_React_Components.docx",
            "React user interface components"),
        ("05_SOURCE_Volume_03_Libraries_and_Contexts.docx",
            "Shared libraries, React contexts, hooks, and configuration"),
        ("06_SOURCE_Volume_04_Database_and_Schema.docx",
            "Prisma schema, migrations, and database access layer"),
        ("07_SOURCE_Volume_05_Configs_and_Documentation.docx",
            "Build configuration, scripts, and technical documentation"),
        ("onekof-platform-source.zip",
            "Complete repository archive (backup copy)"),
        ("README.txt",
            "CD-ROM viewing and extraction instructions"),
    ]

    table = doc.add_table(rows=len(deposit_rows), cols=2)
    table.autofit = False
    table.style = 'Light List Accent 1'
    for i, (name, desc) in enumerate(deposit_rows):
        name_cell = table.rows[i].cells[0]
        desc_cell = table.rows[i].cells[1]
        name_cell.width = Inches(3.2)
        desc_cell.width = Inches(3.3)
        _shade_cell(name_cell, 'E8F5F3')

        np_ = name_cell.paragraphs[0]
        nrun = np_.add_run(name)
        nrun.bold = True
        nrun.font.size = Pt(9)
        nrun.font.name = 'Consolas'
        nrun.font.color.rgb = DARK

        dp = desc_cell.paragraphs[0]
        drun = dp.add_run(desc)
        drun.font.size = Pt(10)
    doc.add_paragraph()

    add_section_heading(doc, "Compliance Note")
    add_paragraph(
        doc,
        "The Onekof source code significantly exceeds the 50-page threshold "
        "referenced in EIPA Form C-1 Section 5. In lieu of the printed "
        "first-25 / last-25 page excerpt, the complete source code is "
        "provided electronically on CD-ROM as permitted for voluminous "
        "software deposits. The deposit is organized into five logical "
        "volumes (Volumes 01\u201305) totaling 132,173 lines of original "
        "source code authored by Oli Tamrat Oli.",
    )

    add_section_heading(doc, "Exclusions from Deposit")
    add_paragraph(
        doc,
        "The following categories are intentionally excluded from the "
        "deposit because they are not original works of the applicant "
        "and/or contain sensitive information:",
    )
    add_bullets(doc, [
        "Third-party dependencies (node_modules, vendored libraries)",
        "Build artifacts (.next, dist, build directories)",
        "Lock files (package-lock.json, pnpm-lock.yaml)",
        "Environment files (.env) containing secrets and credentials",
        "API keys, database passwords, and OAuth client secrets",
        "Four of five compiled locale JSON files (only en.json is retained "
        "as representative; AM/OM/TI/SO translation files are excluded for "
        "deposit volume management)",
    ])

    add_section_heading(doc, "Verification")
    add_paragraph(
        doc,
        "I, Oli Tamrat Oli, certify that the source code contained on the "
        "accompanying CD-ROM is a true and accurate copy of the Onekof "
        "Platform source code as of the deposit generation date above, and "
        "that I am the sole author and copyright holder of the deposited "
        "work.",
    )
    add_field_table(doc, [
        ("Declarant", "Oli Tamrat Oli"),
        ("Signature", ""),
        ("Date", ""),
        ("Place", "Burtonsville, Maryland, USA"),
    ])

    doc.save(output_path)
    print(f"  OK  {output_path.name}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    FORMS_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Generating EIPA forms into: {FORMS_DIR}")
    print("-" * 60)

    generate_form_c1(FORMS_DIR / "Form_C-1_Application.docx")
    generate_software_description(FORMS_DIR / "Software_Description.docx")
    generate_publication_date(FORMS_DIR / "Publication_Date.docx")
    generate_deposit_cover(FORMS_DIR / "Source_Code_Deposit_Cover.docx")

    print("-" * 60)
    print("Done. 4 forms generated.")


if __name__ == '__main__':
    main()
