"""
Generates GO_LIVE_ROADMAP.docx from the roadmap content.
Run: python scripts/generate_roadmap_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ────────────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x1A, 0x56, 0xDB)   # Onekof primary
BRAND_DARK   = RGBColor(0x11, 0x18, 0x27)   # near-black headers
ACCENT_GREEN = RGBColor(0x05, 0x96, 0x69)   # done / success
ACCENT_AMBER = RGBColor(0xD9, 0x77, 0x06)   # pending / warning
ACCENT_RED   = RGBColor(0xDC, 0x26, 0x26)   # risk / critical
TABLE_HEADER = RGBColor(0x1E, 0x40, 0xAF)   # table header bg (deep blue)
TABLE_ALT    = RGBColor(0xEF, 0xF6, 0xFF)   # alternating row bg
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, edges=('top','bottom','left','right')):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in edges:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:color'), 'BFDBFE')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_BLUE
    # bottom border line
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:color'), '1A56DB')
    pBdr.append(bot)
    pPr.append(pBdr)
    return para


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_DARK
    return para


def h3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = BRAND_BLUE
    return para


def body(doc, text, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Inches(indent * 0.25)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return para


def bullet(doc, text, level=0, done=False, pending=False):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    para  = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(2)
    color = ACCENT_GREEN if done else (ACCENT_AMBER if pending else RGBColor(0x1F,0x29,0x37))
    run  = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(11)
    run.font.color.rgb = color
    return para


def code_block(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.left_indent = Inches(0.3)
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F1F5F9')
        para._p.get_or_add_pPr().append(shd)


def status_badge(text):
    if 'DONE' in text or 'Done' in text:
        return f'[DONE] {text.replace("[DONE]","").replace("Done","").strip()}'
    return text


def make_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        set_cell_bg(hdr_cells[i], TABLE_HEADER)
        set_cell_border(hdr_cells[i])
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], TABLE_ALT)
            set_cell_border(cells[c_idx])
            run = cells[c_idx].paragraphs[0].add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            # Color status cells
            v = str(val)
            if 'DONE' in v or 'Done' in v:
                run.font.color.rgb = ACCENT_GREEN
                run.bold = True
            elif 'Pending' in v or 'Blocked' in v:
                run.font.color.rgb = ACCENT_AMBER

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ── Cover / Title ─────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(24)
title_para.paragraph_format.space_after  = Pt(6)
r = title_para.add_run('Onekof PM')
r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(32)
r.font.color.rgb = BRAND_BLUE

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_after = Pt(4)
r = sub_para.add_run('Deployment, Migration & Go-Live Roadmap')
r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(18)
r.font.color.rgb = BRAND_DARK

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(2)
add_run(meta, 'Author: Oli T. Oli / DAPS Analytics     |     Date: 2026-05-23     |     Classification: Internal / Confidential',
        size=10, color=RGBColor(0x6B, 0x72, 0x80))

doc.add_paragraph()

# ── Executive Summary ──────────────────────────────────────────────────────
h1(doc, 'Executive Summary')

make_table(doc,
    headers=['Milestone', 'Target Date', 'Status'],
    rows=[
        ['Tier 3 CI/CD live (Vercel + Supabase)',       '2026-05-23', 'DONE'],
        ['Wave 5 INSA security hardening (P1–P5)',       '2026-05-23', 'DONE'],
        ['EthioTelecom VM provisioned',                  '2026-06-07', 'Pending'],
        ['onekof.et domain registered',                  '2026-06-07', 'Pending'],
        ['Tier 2 go-live (Ethiopia self-hosted)',         '2026-06-14', 'Pending'],
        ['Internal beta launch',                         '2026-06-21', 'Pending'],
        ['INSA certification submitted',                 '2026-06-28', 'Pending'],
        ['External / public beta',                       '2026-07-05', 'Pending'],
        ['First government pilot outreach',              '2026-07-12', 'Pending'],
        ['Tier 1 sovereign pilot deployment',            '2026-08-01', 'Pending'],
    ],
    col_widths=[3.2, 1.5, 1.5]
)

# ── Part 1 — Current State ─────────────────────────────────────────────────
h1(doc, 'Part 1 — Current State (as of 2026-05-23)')

h2(doc, 'Completed Infrastructure')
make_table(doc,
    headers=['Component', 'State'],
    rows=[
        ['Next.js 14 web app',              'Built — standalone Docker image ~300 MB'],
        ['Prisma 5 + PostgreSQL 15',         '14 migrations applied on production DB'],
        ['NextAuth.js v5',                   'Working — credentials + session management'],
        ['CI workflow (GitHub Actions)',      'Green — build + type check on every push'],
        ['Deploy workflow (GitHub Actions)', 'Green — auto-deploys to Vercel on master push'],
        ['docker-compose.prod.yml',          'Ready for Tier 2'],
        ['Caddyfile (auto-SSL)',             'Ready for Tier 2'],
        ['setup-tier2-server.sh',            'Idempotent setup script ready'],
        ['.env.production.example',          'Complete with all required keys'],
        ['Mobile app (Expo/RN)',             'Feature-complete, EAS build linked'],
        ['INSA security P1–P6',             'DONE — Wave 5 closed today'],
        ['EIPA copyright registration',      'Package submitted 2026-04-11'],
    ],
    col_widths=[2.5, 4.2]
)

h2(doc, 'Remaining External Blockers')
make_table(doc,
    headers=['Item', 'Blocker', 'Owner'],
    rows=[
        ['onekof.et domain',         'EthioTelecom/nic.et registration',    'Oli'],
        ['EthioTelecom VM',          'Server procurement + SSH access',      'Oli'],
        ['Resend domain verify',     'Needs .et domain first',               'Oli'],
        ['INSA cert submission',     'Ready to submit — all gaps closed',    'Oli'],
    ],
    col_widths=[2.0, 3.2, 1.5]
)

# ── Part 2 — Architecture ──────────────────────────────────────────────────
h1(doc, 'Part 2 — Three-Tier Deployment Architecture')

tiers = [
    ('Tier 3 — Vercel + Supabase', 'LIVE',
     'URL: onekof.vercel.app → onekof.com\nDB: Supabase PostgreSQL + Session Pooler\n'
     'Redis: Upstash\nDeploy: GitHub push to master → auto-deploy\n'
     'Users: Global cloud SaaS, diaspora, international orgs'),
    ('Tier 2 — Self-Hosted Ethiopia', 'Target: 2026-06-14',
     'URL: onekof.et\nServer: EthioTelecom VM (Ubuntu 22.04, 4 vCPU, 8 GB RAM)\n'
     'DB: PostgreSQL 15 container\nRedis: Redis 7 container\nSSL: Caddy auto Let\'s Encrypt\n'
     'Users: Ethiopian businesses preferring data residency'),
    ('Tier 1 — Air-Gapped Sovereign', 'Target: 2026-08-01',
     'URL: Internal — customer\'s own network\nDelivery: Docker image via USB or ACR pull-token\n'
     'DB: Customer-managed PostgreSQL\nAuth: Fully isolated — no internet required\n'
     'Users: Ethiopian government, INSA-classified environments'),
]

for name, status, desc in tiers:
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells

    set_cell_bg(cells[0], TABLE_HEADER)
    cells[0].width = Inches(2.2)
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run(name)
    r0.bold = True; r0.font.name = 'Calibri'; r0.font.size = Pt(11)
    r0.font.color.rgb = WHITE
    p0.add_run('\n')
    r1 = p0.add_run(status)
    r1.font.name = 'Calibri'; r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)

    cells[1].width = Inches(4.5)
    p1 = cells[1].paragraphs[0]
    r2 = p1.add_run(desc)
    r2.font.name = 'Calibri'; r2.font.size = Pt(10)
    r2.font.color.rgb = BRAND_DARK

    doc.add_paragraph()

# ── Part 3 — Week-by-Week Plan ─────────────────────────────────────────────
h1(doc, 'Part 3 — Week-by-Week Execution Plan')

weeks = [
    ('Week 1: May 25–31', 'Domain + Server Prep', [
        ('Register onekof.et at nic.et / EthioTelecom (bring: business cert, TIN)', False),
        ('Contact EthioTelecom for VM quote: 4 vCPU / 8 GB / 100 GB SSD / Ubuntu 22.04', False),
        ('Verify Vercel env vars (vercel env ls) — confirm all required keys present', False),
        ('Tag current commit: git tag v1.0.0-beta && git push origin v1.0.0-beta', False),
        ('Compile INSA certification submission package', False),
        ('Submit INSA package to Technology House (or portal) by May 31', False),
    ]),
    ('Week 2: June 1–7', 'Server Setup + Smoke Test', [
        ('SSH into EthioTelecom VM, harden SSH (disable password auth), configure UFW', False),
        ('Clone repo: git clone https://github.com/daps-analytics/onekof-platform.git /opt/onekof', False),
        ('Run: bash scripts/setup-tier2-server.sh (installs Docker, generates secrets, starts stack)', False),
        ('Edit .env.production: fill in NEXTAUTH_URL, RESEND_API_KEY, SENTRY_DSN', False),
        ('Tier 2 smoke test: all services healthy, signup + login + project creation works', False),
        ('Verify blob encryption: uploaded files unreadable on disk (xxd check)', False),
    ]),
    ('Week 3: June 8–14', 'DNS Cutover + Tier 2 Go-Live', [
        ('Point DNS A record: onekof.et → <server-ip>', False),
        ('Point DNS A record: www.onekof.et → <server-ip>', False),
        ('Add Resend SPF + DKIM + DMARC records, verify domain', False),
        ('Restart Caddy → auto Let\'s Encrypt cert issued (~30 seconds)', False),
        ('SSL Labs check: https://ssllabs.com → target A+ rating', False),
        ('Security headers check: https://securityheaders.com → target A', False),
        ('Test email delivery end-to-end (signup verification email)', False),
        ('Take final Supabase backup before migration starts', False),
        ('TIER 2 GO-LIVE DECLARATION — June 14', False),
    ]),
    ('Week 4: June 15–21', 'Internal Beta', [
        ('Create beta organization, invite 5–10 internal / trusted users', False),
        ('Configure Sentry alerts: any P0/P1 error → email + Telegram notification', False),
        ('Configure UptimeRobot: downtime > 2 min → SMS alert', False),
        ('Run k6 load test: 50 concurrent users, 5 min — target p95 < 500 ms', False),
        ('Fix TypeScript errors: change ignoreBuildErrors to false, run pnpm dedupe', False),
        ('Daily monitoring: docker compose ps, Sentry dashboard, disk usage', False),
    ]),
    ('Week 5: June 22–28', 'INSA Follow-Up + Public Beta Prep', [
        ('Follow up on INSA submission (submitted May 30)', False),
        ('Prepare 1-page data residency letter (Amharic + English)', False),
        ('Update onekof.et landing page: signup CTA + pricing page', False),
        ('Configure support@onekof.et email forwarding', False),
        ('Build EAS production APK: eas build --platform android --profile production', False),
        ('Distribute Android APK via Telegram channel + onekof.et/download', False),
    ]),
    ('Week 6: June 29 – July 5', 'Public Beta Launch', [
        ('Final pre-launch checklist: load test pass, SSL A+, Sentry active, uptime monitor active', False),
        ('Announce on LinkedIn in English + Amharic', False),
        ('Post in Ethiopian tech Telegram groups', False),
        ('Send email to beta waitlist', False),
        ('Monitor Sentry continuously for first 4 hours', False),
        ('Respond to all support inquiries within 4 hours', False),
    ]),
    ('Week 7–8: July 6–19', 'Government Pilot Outreach', [
        ('Priority 1: Ministry of Innovation & Technology (MInT) — Director-General\'s office', False),
        ('Priority 2: Ethiopian Investment Commission — IT Director', False),
        ('Priority 3: Commercial Bank of Ethiopia — Digital Banking division', False),
        ('Deliver: 1-page executive summary (Amharic + English), INSA cert, data residency letter', False),
        ('Set up government demo environment with sample projects', False),
        ('Present Tier 1 deployment brief: runs fully offline in their data center', False),
    ]),
    ('Week 9–12: July 20 – August 15', 'Tier 1 Sovereign Deploy', [
        ('Tag release: git tag v1.0.0-tier1-20260720 && git push origin --tags', False),
        ('Pull and save Docker image: docker save ... | gzip > onekof-tier1-v1.0.0.tar.gz', False),
        ('Generate SHA256 hash for integrity verification', False),
        ('Prepare USB delivery package: image + compose + env template + runbook (Amharic)', False),
        ('On-site or remote: docker load, fill .env.production, docker compose up, run migrations', False),
        ('TIER 1 GO-LIVE: first sovereign government deployment', False),
    ]),
]

for week_title, week_goal, tasks in weeks:
    h2(doc, week_title)
    body(doc, f'Goal: {week_goal}')
    for task, done in tasks:
        bullet(doc, ('✓  ' if done else '☐  ') + task, done=done)
    doc.add_paragraph()

# ── Part 4 — Data Migration ────────────────────────────────────────────────
h1(doc, 'Part 4 — Data Migration Plan (Tier 3 Cloud → Tier 2 Self-Hosted)')

body(doc, 'Migration window: approximately 2 hours. Each organization\'s data is isolated — migrating one org does not affect others.')

h2(doc, 'Migration Steps')

migration_steps = [
    ('T-48 hours', 'Pre-migration backup from Supabase using pg_dump --format=custom'),
    ('T-24 hours', 'Dry run restore on staging Tier 2 — verify record counts match source'),
    ('T-60 min',   'Notify affected users of maintenance window'),
    ('T-30 min',   'Set NEXT_PUBLIC_MAINTENANCE=true → deploy (shows maintenance page)'),
    ('T-0',        'Final export from Supabase (captures last 30 min of data)'),
    ('T+10 min',   'pg_restore into Tier 2 PostgreSQL container'),
    ('T+15 min',   'Verify data counts match source exactly'),
    ('T+20 min',   'Update DNS / give users new URL (onekof.et/[org-slug])'),
    ('T+25 min',   'Remove maintenance mode — NEXT_PUBLIC_MAINTENANCE=false → deploy'),
    ('T+30 min',   'Org admin acceptance test — verify their data'),
    ('T+60 min',   'Decommission org data from Supabase (data residency compliance)'),
]

make_table(doc,
    headers=['Time', 'Action'],
    rows=migration_steps,
    col_widths=[1.2, 5.5]
)

# ── Part 5 — Rollback ──────────────────────────────────────────────────────
h1(doc, 'Part 5 — Rollback Plans')

h2(doc, 'Tier 3 — Vercel (Instant)')
code_block(doc, [
    'vercel rollback',
    '',
    '# Or via dashboard:',
    '# Vercel → Project → Deployments → [previous] → ... → Promote to Production',
])

h2(doc, 'Tier 2 — Self-Hosted')
code_block(doc, [
    '# 1. Stop app container',
    'docker compose -f docker-compose.prod.yml stop onekof-web',
    '',
    '# 2. Set previous image tag in .env.production',
    'nano .env.production   # DOCKER_IMAGE=ghcr.io/daps-analytics/onekof-web:vPREVIOUS',
    '',
    '# 3. Pull and restart',
    'docker compose -f docker-compose.prod.yml pull onekof-web',
    'docker compose -f docker-compose.prod.yml up -d onekof-web',
])

h2(doc, 'Tier 1 — Sovereign')
body(doc, 'Load previous USB image: docker load < onekof-tier1-vPREVIOUS.tar.gz')
body(doc, 'Customer IT team restores from their local backup (procedure included in Tier 1 runbook).')

# ── Part 6 — Success Metrics ───────────────────────────────────────────────
h1(doc, 'Part 6 — Success Metrics')

h2(doc, 'Beta Phase (Weeks 4–6)')
make_table(doc,
    headers=['Metric', 'Target'],
    rows=[
        ['Internal beta users',        '10+'],
        ['Projects created',           '5+'],
        ['Issues created',             '50+'],
        ['P0 production incidents',    '0'],
        ['Uptime',                     '> 99.5%'],
        ['Page load (p95)',            '< 2 seconds'],
        ['API response (p95)',         '< 500 ms'],
    ],
    col_widths=[3.0, 3.7]
)

h2(doc, 'Public Beta Phase (Weeks 6–8)')
make_table(doc,
    headers=['Metric', 'Target'],
    rows=[
        ['Registered users',           '50+'],
        ['Active organizations',       '10+'],
        ['Weekly active users',        '25+'],
        ['Email delivery rate',        '> 98%'],
        ['P1 bug reports',             '< 5'],
    ],
    col_widths=[3.0, 3.7]
)

h2(doc, 'Government Pilot Phase (Weeks 9–12)')
make_table(doc,
    headers=['Metric', 'Target'],
    rows=[
        ['Government orgs in pilot',       '1–2'],
        ['Tier 1 deployments',             '1'],
        ['INSA certification status',      'Submitted (awaiting response)'],
        ['Enterprise contract conversations', '3+ active'],
    ],
    col_widths=[3.0, 3.7]
)

# ── Part 7 — Open Items ────────────────────────────────────────────────────
h1(doc, 'Part 7 — Open Items Tracker')

make_table(doc,
    headers=['#', 'Item', 'Owner', 'Target', 'Status'],
    rows=[
        ['OI-1',  'Register onekof.et domain',                    'Oli', 'Jun 7',  'Pending — do first'],
        ['OI-2',  'GitHub Actions CI/CD workflows',               'Oli', 'Done',   'Done'],
        ['OI-3',  'docker-compose.prod.yml',                      'Oli', 'Done',   'Done'],
        ['OI-4',  'Caddyfile (auto-SSL)',                         'Oli', 'Done',   'Done'],
        ['OI-5',  'GitHub repository secrets',                    'Oli', 'Done',   'Done'],
        ['OI-6',  'INSA P1–P5 security gaps (Wave 5)',            'Oli', 'Done',   'Done'],
        ['OI-7',  'Provision EthioTelecom VM',                    'Oli', 'Jun 7',  'Pending'],
        ['OI-8',  'Run setup-tier2-server.sh',                    'Oli', 'Jun 7',  'Blocked on OI-7'],
        ['OI-9',  'Tier 2 smoke test',                            'Oli', 'Jun 14', 'Blocked on OI-7'],
        ['OI-10', 'Resend domain + email test',                   'Oli', 'Jun 14', 'Blocked on OI-1'],
        ['OI-11', 'INSA certification submission',                'Oli', 'May 31', 'Pending — this week'],
        ['OI-12', 'Remove ignoreBuildErrors: true',               'Oli', 'Jun 21', 'Beta week'],
        ['OI-13', 'Load test (k6)',                               'Oli', 'Jun 21', 'Beta week'],
        ['OI-14', 'Mobile app store submission',                  'Oli', 'Jun 28', 'Pre-launch'],
        ['OI-15', 'Tier 1 offline USB package',                   'Oli', 'Jul 20', 'Gov pilot prep'],
        ['OI-16', 'Amharic ops runbook',                          'Oli', 'Jul 20', 'Gov pilot prep'],
        ['OI-17', 'Government pilot outreach',                    'Oli', 'Jul 12', 'Post-public-beta'],
        ['OI-18', 'First Tier 1 sovereign deployment',            'Oli', 'Aug 1',  'Post-INSA cert'],
    ],
    col_widths=[0.5, 2.8, 0.6, 0.8, 1.5]
)

# ── Part 8 — Risk Register ─────────────────────────────────────────────────
h1(doc, 'Part 8 — Risk Register')

make_table(doc,
    headers=['Risk', 'Probability', 'Impact', 'Mitigation'],
    rows=[
        ['EthioTelecom VM procurement delays',
         'Medium', 'High',
         'Use DigitalOcean/Linode as interim; migrate when ET VM ready'],
        ['.et domain registration > 2 weeks',
         'Medium', 'Medium',
         'Launch on onekof.com interim; cut over to .et when ready'],
        ['INSA certification takes > 6 weeks',
         'Low', 'High',
         'Start gov outreach in parallel — cert strengthens, not blocks'],
        ['Blob encryption breaks existing files on upgrade',
         'Low', 'High',
         'decryptBlobBuffer is backward-compatible — dev passthrough for unencrypted files'],
        ['Load spikes overwhelm single Tier 2 VM',
         'Low', 'Medium',
         'Caddy + Next.js handle ~500 concurrent on 4 vCPU; add scale if needed'],
        ['Government procurement cycle longer than expected',
         'High', 'Medium',
         'Use time to onboard commercial customers — gov revenue is upside, not only path'],
    ],
    col_widths=[2.0, 0.9, 0.7, 3.1]
)

# ── Part 9 — Immediate Actions ─────────────────────────────────────────────
h1(doc, 'Part 9 — Immediate Actions This Week (May 25–31)')

body(doc, 'These three actions unblock everything else. Do them in order.')

actions = [
    ('Action 1 — Register onekof.et (May 25)',
     'Go to nic.et or EthioTelecom domain services in person.\n'
     'Bring: Commercial registration certificate, TIN certificate, technical contact email.\n'
     'Processing time: 3–7 business days. Start immediately.'),
    ('Action 2 — Submit INSA Certification (May 26–27)',
     'All P1–P5 security gaps are now closed. You are ready to submit.\n'
     'Package: GO_LIVE_ROADMAP.md + MIGRATION_DEPLOYMENT_PLAN.md + Wave 4/5 security audit summary + EIPA confirmation.\n'
     'Submit to INSA at Technology House or via their official portal.'),
    ('Action 3 — Procurement Call with EthioTelecom (May 27–28)',
     'Request: 1x Ubuntu 22.04 VM, 4 vCPU / 8 GB RAM / 100 GB SSD, static public IPv4.\n'
     'Ports: 80 and 443 open inbound/outbound. Port 22 restricted to your IP.\n'
     'Estimated cost: 3,000–8,000 ETB/month. Negotiate annual contract.'),
]

for title, desc in actions:
    h2(doc, title)
    for line in desc.split('\n'):
        body(doc, line)

# ── Footer ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer_para,
        'Document Owner: Oli T. Oli / DAPS Analytics   |   Next Review: 2026-06-14   |   Source of Truth: GO_LIVE_ROADMAP.md',
        size=9, color=RGBColor(0x9C, 0xA3, 0xAF))

# ── Save ───────────────────────────────────────────────────────────────────
output_path = 'GO_LIVE_ROADMAP.docx'
doc.save(output_path)
print(f'Done. Saved to: {output_path}')
