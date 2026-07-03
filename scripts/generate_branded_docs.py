"""
Onekof PM — Branded Document Generator
Generates two .docx files with full Onekof brand styling:
  1. ONEKOF_PROJECT_READINESS_REPORT.docx
  2. ONEKOF_MIGRATION_DEPLOYMENT_PLAN.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from datetime import date

# ─── Brand Colors ────────────────────────────────────────────────────────────
TEAL        = RGBColor(0x1C, 0x8C, 0x7D)   # #1C8C7D primary
TEAL_LIGHT  = RGBColor(0x2B, 0xB5, 0xA2)   # #2BB5A2 teal-light
DARK_BG     = RGBColor(0x0B, 0x0E, 0x11)   # #0B0E11 page bg
SLATE       = RGBColor(0x1E, 0x29, 0x3B)   # section heading bg
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = RGBColor(0xF8, 0xFA, 0xFC)
MID_GRAY    = RGBColor(0xE2, 0xE8, 0xF0)
TEXT_DARK   = RGBColor(0x0F, 0x17, 0x2A)
TEXT_MID    = RGBColor(0x47, 0x55, 0x69)
GREEN_DARK  = RGBColor(0x06, 0x6B, 0x4F)
GREEN_BG    = RGBColor(0xD1, 0xFA, 0xE5)
AMBER_DARK  = RGBColor(0x92, 0x40, 0x0E)
AMBER_BG    = RGBColor(0xFF, 0xED, 0xC2)
RED_DARK    = RGBColor(0x99, 0x1B, 0x1B)
RED_BG      = RGBColor(0xFE, 0xE2, 0xE2)
BLUE_DARK   = RGBColor(0x1E, 0x40, 0xAF)
BLUE_BG     = RGBColor(0xDB, 0xEA, 0xFE)
VIOLET_DARK = RGBColor(0x5B, 0x21, 0xB6)
VIOLET_BG   = RGBColor(0xED, 0xE9, 0xFE)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_para_border_bottom(para, color="1C8C7D", size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_page_margins(doc, top=1.8, bottom=1.8, left=2.0, right=2.0):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font='Calibri'):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def add_para(doc, text='', bold=False, italic=False, size=11,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=4, font='Calibri'):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if text:
        add_run(para, text, bold=bold, italic=italic,
                size=size, color=color, font=font)
    return para

def section_heading(doc, text, level=1):
    """Teal section heading with bottom border."""
    sizes = {1: 16, 2: 13, 3: 11}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)
    if level == 1:
        set_para_border_bottom(para, color="1C8C7D", size=16)
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(sizes.get(level, 11))
    run.font.bold  = True
    run.font.color.rgb = TEAL if level <= 2 else TEXT_DARK
    return para

def status_badge(text):
    """Return (display_text, bg_color, text_color) tuple."""
    t = text.strip().upper()
    if any(x in t for x in ['✅', 'DONE', 'COMPLETE', 'LIVE', 'ACTIVE', 'CLOSED', 'GREEN']):
        return (text, GREEN_BG, GREEN_DARK)
    if any(x in t for x in ['⚠️', 'PARTIAL', 'IN PROGRESS', 'PENDING']):
        return (text, AMBER_BG, AMBER_DARK)
    if any(x in t for x in ['❌', 'BLOCKER', 'NOT STARTED', 'MISSING']):
        return (text, RED_BG, RED_DARK)
    if any(x in t for x in ['⏳', 'PLANNED', 'DEFERRED', 'TBD']):
        return (text, BLUE_BG, BLUE_DARK)
    return (text, LIGHT_GRAY, TEXT_DARK)

def add_table(doc, headers, rows, col_widths=None, header_bg=TEAL,
              stripe=True, status_col=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        run = para.add_run(h)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(9)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = LIGHT_GRAY if (stripe and r_idx % 2 == 0) else WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

            # Status column gets color badge treatment
            if status_col is not None and c_idx == status_col:
                _, s_bg, s_fg = status_badge(str(val))
                set_cell_bg(cell, s_bg)
                run = para.add_run(str(val))
                run.font.name  = 'Calibri'
                run.font.size  = Pt(9)
                run.font.bold  = True
                run.font.color.rgb = s_fg
            else:
                set_cell_bg(cell, bg)
                run = para.add_run(str(val))
                run.font.name  = 'Calibri'
                run.font.size  = Pt(9)
                run.font.color.rgb = TEXT_DARK

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_cover_page(doc, title, subtitle, date_str, classification='Confidential'):
    """Full branded cover page."""
    # Top teal bar via shaded paragraph
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    bar_run = bar.add_run(' ' * 200)
    bar_run.font.size = Pt(2)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1C8C7D')
    bar._p.get_or_add_pPr().append(shd)

    doc.add_paragraph().paragraph_format.space_after = Pt(60)

    # ONEKOF wordmark
    wm = doc.add_paragraph()
    wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wm.paragraph_format.space_after = Pt(4)
    r1 = wm.add_run('ONEKOF')
    r1.font.name  = 'Calibri'
    r1.font.size  = Pt(36)
    r1.font.bold  = True
    r1.font.color.rgb = TEAL

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(40)
    tl = tagline.add_run('Where African teams do their best work')
    tl.font.name   = 'Calibri'
    tl.font.size   = Pt(11)
    tl.font.italic = True
    tl.font.color.rgb = TEXT_MID

    # Teal divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_after = Pt(30)
    set_para_border_bottom(div, color='1C8C7D', size=24)

    # Document title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(10)
    tr = t.add_run(title)
    tr.font.name  = 'Calibri'
    tr.font.size  = Pt(26)
    tr.font.bold  = True
    tr.font.color.rgb = TEXT_DARK

    # Subtitle
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(60)
    str_ = st.add_run(subtitle)
    str_.font.name  = 'Calibri'
    str_.font.size  = Pt(13)
    str_.font.color.rgb = TEXT_MID

    # Meta block
    for label, value in [
        ('Date', date_str),
        ('Author', 'Oli T. Oli — Founder & Lead Engineer'),
        ('Company', 'DAPS Analytics'),
        ('Classification', classification),
    ]:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(3)
        ml = meta.add_run(f'{label}:  ')
        ml.font.name  = 'Calibri'
        ml.font.size  = Pt(10)
        ml.font.bold  = True
        ml.font.color.rgb = TEAL
        mv = meta.add_run(value)
        mv.font.name  = 'Calibri'
        mv.font.size  = Pt(10)
        mv.font.color.rgb = TEXT_DARK

    doc.add_page_break()

def add_highlight_box(doc, text, bg=TEAL, fg=WHITE, size=11):
    """Full-width shaded callout paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.bold  = True
    run.font.color.rgb = fg
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_bg = f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}"
    shd.set(qn('w:fill'), hex_bg)
    para._p.get_or_add_pPr().append(shd)
    return para

def bullet(doc, text, level=0, color=TEXT_DARK):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    run.font.color.rgb = color
    return para

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — PROJECT READINESS REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def build_readiness_report():
    doc = Document()
    set_page_margins(doc)

    # Default paragraph font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # Cover
    add_cover_page(
        doc,
        title='Project Readiness Report',
        subtitle='Founders & Stakeholders Briefing — May 2026',
        date_str='May 23, 2026',
        classification='Confidential — Internal Distribution Only'
    )

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    section_heading(doc, '1.  Executive Summary')
    add_highlight_box(
        doc,
        '  Onekof PM has achieved full technical production readiness. '
        'All engineering blockers are resolved. '
        'The only remaining go-live steps are founder actions: domain, server, and INSA certification submission.',
        bg=TEAL, fg=WHITE, size=11
    )
    add_para(doc, space_after=6)
    add_para(doc,
        'Onekof PM is an enterprise-grade, multi-tenant project management platform built '
        'specifically for Ethiopian and East African organizations. As of May 23, 2026, the platform '
        'passes 100% of INSA security requirements, enforces zero TypeScript build errors, and '
        'operates with full error tracking, audit logging, and five-language support.',
        size=10, color=TEXT_DARK, space_after=6)
    add_para(doc,
        'The platform is live on global cloud infrastructure (Tier 3) and the same Docker image '
        'is ready to deploy on an EthioTelecom server (Tier 2) the moment the domain and VM are provisioned.',
        size=10, color=TEXT_DARK, space_after=10)

    # KPI strip table
    add_table(doc,
        headers=['Metric', 'Value'],
        rows=[
            ['TypeScript errors', '0  —  Strict mode enforced'],
            ['INSA security compliance', '100%  —  All P1–P6 gaps closed'],
            ['API routes auth-guarded', '141 / 141'],
            ['Languages supported', '5  —  EN, AM, OM, TI, SO'],
            ['Docker image size', '408 MB standalone'],
            ['CI/CD pipelines', 'GitHub Actions — both green'],
            ['Error tracking', 'Sentry live  (web + mobile)'],
            ['Audit log', 'Append-only, 10 routes, INSA-compliant'],
        ],
        col_widths=[2.8, 4.0],
        header_bg=SLATE,
        status_col=1
    )

    # ── 2. Product Overview ──────────────────────────────────────────────────
    section_heading(doc, '2.  Product Overview')
    add_table(doc,
        headers=['Attribute', 'Detail'],
        rows=[
            ['Product name', 'Onekof PM'],
            ['Category', 'Enterprise Project Management SaaS'],
            ['Primary market', 'Ethiopian government, NGOs, private sector (East Africa)'],
            ['IP owner', 'DAPS Analytics  (commercial rights)'],
            ['Author / moral rights', 'Oli Tamrat Oli  (non-transferable, Proclamation 410/2004)'],
            ['IP status', 'EIPA copyright deposit in progress'],
            ['Stage', 'Pre-launch — no paying customers'],
            ['Live URL', 'onekof.com  (Vercel global cloud)'],
        ],
        col_widths=[2.4, 4.4],
        header_bg=TEAL
    )

    section_heading(doc, 'Core Feature Set', level=2)
    add_table(doc,
        headers=['Module', 'Status'],
        rows=[
            ['Projects & Issues (Kanban, Backlog, List)', '✅ Complete'],
            ['Goals & Key Results (OKR framework)', '✅ Complete'],
            ['Teams & Member management', '✅ Complete'],
            ['Budget tracking (ETB, multi-currency)', '✅ Complete'],
            ['Wiki / Knowledge base', '✅ Complete'],
            ['Document management + AI processing', '✅ Complete'],
            ['Audit log (INSA-compliant, append-only)', '✅ Complete'],
            ['Notifications (web + mobile push)', '✅ Complete'],
            ['Automations & rules engine', '✅ Complete'],
            ['Analytics & dashboard', '✅ Complete'],
            ['Mobile app — iOS + Android (Expo/RN)', '✅ Feature-complete'],
            ['AI features (Anthropic)', '⏳ Ready — pending API key'],
        ],
        col_widths=[4.6, 2.2],
        header_bg=TEAL,
        status_col=1
    )

    # ── 3. Technical Readiness ───────────────────────────────────────────────
    section_heading(doc, '3.  Technical Readiness')

    section_heading(doc, 'INSA Security — 100% Complete', level=2)
    add_table(doc,
        headers=['#', 'Requirement', 'Delivered', 'Status'],
        rows=[
            ['P1', 'CSRF origin validation on all mutation APIs', 'middleware.ts enforceCsrfOrigin()', '✅ Closed'],
            ['P2', 'Admin endpoint rate limiting (60 req/min/IP)', 'middleware.ts checkAdminRateLimit()', '✅ Closed'],
            ['P3', 'Audit log immutability — DELETE → 405', 'audit-log/route.ts explicit 405', '✅ Closed'],
            ['P4', 'AES-256-GCM at-rest encryption for blobs', 'local-fs.ts + BLOB_ENCRYPTION_KEY in Vercel', '✅ Closed'],
            ['P5', 'Session invalidation on password change', 'change-password route → invalidateAllUserSessions()', '✅ Closed'],
            ['P6', 'HTTP security headers', 'CSP, X-Frame, HSTS, Referrer-Policy in middleware', '✅ Closed'],
        ],
        col_widths=[0.35, 2.2, 2.8, 1.1],
        header_bg=TEAL,
        status_col=3
    )

    section_heading(doc, 'Infrastructure Readiness by Tier', level=2)
    add_table(doc,
        headers=['Component', 'Tier 3 — Cloud', 'Tier 2 — Self-Hosted', 'Tier 1 — Sovereign'],
        rows=[
            ['Web app', '✅ Vercel (fra1)', '✅ Docker 408 MB', '✅ Same image'],
            ['Database', '✅ Supabase PG15', '✅ PostgreSQL 15', '✅ Same config'],
            ['Redis', '✅ Upstash', '✅ Redis 7', '✅ Same config'],
            ['File storage', '✅ Vercel Blob', '✅ Local-fs + AES-256', '✅ Same driver'],
            ['SSL', '✅ Vercel auto', '✅ Caddy auto', 'Manual'],
            ['Backups', '✅ Automated', '✅ GPG cron', '✅ Same script'],
            ['CI/CD', '✅ GitHub Actions', 'Manual trigger', 'USB delivery'],
            ['Error tracking', '✅ Sentry live', '⏳ After Tier 2 live', '⏳ After Tier 2'],
        ],
        col_widths=[1.8, 1.8, 1.8, 1.5],
        header_bg=SLATE,
        status_col=1
    )

    # ── 4. Go-Live Readiness Scorecard ───────────────────────────────────────
    section_heading(doc, '4.  Go-Live Readiness Scorecard')

    section_heading(doc, 'Engineering (completed)', level=2)
    add_table(doc,
        headers=['Area', 'Score', 'Notes'],
        rows=[
            ['Core product features', '10 / 10', 'All modules complete'],
            ['Security compliance (INSA)', '10 / 10', 'P1–P6 all closed — Wave 5'],
            ['TypeScript build quality', '10 / 10', 'Zero errors, strict mode'],
            ['Mobile app', '9 / 10', 'Feature-complete, App Store in progress'],
            ['Error tracking (Sentry)', '10 / 10', 'Web + mobile projects live'],
            ['i18n — 5 languages', '10 / 10', 'Fully aligned across all locales'],
            ['CI/CD pipeline', '10 / 10', 'GitHub Actions live and green'],
            ['Docker / self-hosted', '10 / 10', '408 MB image, proven locally'],
            ['Documentation', '9 / 10', 'Runbooks + deployment plan complete'],
        ],
        col_widths=[2.8, 1.2, 2.8],
        header_bg=GREEN_DARK,
        status_col=1
    )

    add_highlight_box(doc, '  Engineering Overall:  98 / 100', bg=GREEN_DARK, fg=WHITE)
    add_para(doc, space_after=6)

    section_heading(doc, 'Business Actions Required (DAPS Analytics)', level=2)
    add_table(doc,
        headers=['Item', 'Score', 'Owner / Status'],
        rows=[
            ['EIPA copyright registration', '3 / 10', 'DAPS Analytics — filing in progress ✅ — Jun 2026'],
            ['INSA certification submission', '2 / 10', 'DAPS Analytics — code ready, submit Jun 2026  ⚡'],
            ['onekof.et domain — EthioTelecom', '1 / 10', 'DAPS Analytics — negotiate Jul 2026 (after EIPA)'],
            ['EthioTelecom VM provisioning', '1 / 10', 'DAPS Analytics — provision Jul–Aug 2026'],
            ['App Store (iOS)', '7 / 10', 'DAPS Analytics — Apple Dev active ✅ — in progress'],
            ['Play Store (Android)', '7 / 10', 'DAPS Analytics — Google Play active ✅ — in progress'],
            ['Cloudflare Full SSL', '2 / 10', 'DAPS Analytics — $10/mo ACM or DNS move to Vercel'],
            ['First pilot customer', '0 / 10', 'DAPS Analytics — outreach to begin Sep 2026'],
        ],
        col_widths=[2.5, 1.0, 3.3],
        header_bg=SLATE,
        status_col=1
    )

    # ── 5. Three-Tier Architecture ───────────────────────────────────────────
    section_heading(doc, '5.  Three-Tier Deployment Architecture')
    add_para(doc,
        'Onekof uses a tiered deployment model so the same Docker image runs across all environments. '
        'Runtime behaviour is controlled entirely by environment variables — no code differences between tiers.',
        size=10, color=TEXT_DARK, space_after=8)

    tiers = [
        ('TIER 3 — Global Cloud  (CURRENT PRODUCTION)',
         '*.onekof.com', TEAL,
         ['Vercel edge network + auto-scaling', 'Supabase PostgreSQL 15', 'Upstash Redis', 'Vercel Blob storage',
          'CI/CD: auto-deploy on every push to master', 'Status: LIVE ✅']),
        ('TIER 2 — Ethiopian Self-Hosted  (NEXT MILESTONE)',
         '*.onekof.et', SLATE,
         ['Docker Compose on EthioTelecom VM', 'Local PostgreSQL 15 + Redis 7', 'AES-256-GCM encrypted file storage',
          'Caddy auto-SSL (Let\'s Encrypt)', 'Automated GPG-encrypted backups', 'Status: CODE-READY — awaiting VM + domain']),
        ('TIER 1 — Government / Air-Gapped  (POST-INSA)',
         'Internal network', RGBColor(0x4B, 0x55, 0x63),
         ['Same Docker image, delivered via USB or ACR pull-token', 'Fully offline — no external DNS required',
          'Designed for INSA-classified environments', 'INSA compliance: 100% ✅', 'Status: CODE-READY — awaiting certification']),
    ]

    for (title, url, color, points) in tiers:
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(8)
        hdr.paragraph_format.space_after  = Pt(2)
        hr = hdr.add_run(f'  {title}')
        hr.font.name  = 'Calibri'
        hr.font.size  = Pt(10)
        hr.font.bold  = True
        hr.font.color.rgb = WHITE
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        hdr._p.get_or_add_pPr().append(shd)

        url_p = doc.add_paragraph()
        url_p.paragraph_format.space_after = Pt(2)
        url_r = url_p.add_run(f'  URL: {url}')
        url_r.font.name  = 'Calibri'
        url_r.font.size  = Pt(9)
        url_r.font.italic = True
        url_r.font.color.rgb = TEXT_MID

        for point in points:
            bullet(doc, point, color=TEXT_DARK)

    add_para(doc, space_after=4)
    add_highlight_box(doc,
        '  Key Principle:  One codebase. One Docker image. '
        'Runtime behaviour controlled entirely by environment variables.',
        bg=TEAL_LIGHT, fg=WHITE)

    # ── 6. Timeline to Launch ────────────────────────────────────────────────
    section_heading(doc, '6.  Timeline to Launch')
    add_para(doc,
        'Timeline is calibrated for Ethiopian bureaucratic realities and the upcoming election cycle. '
        'EIPA registration and INSA certification consume June in full. '
        'Completion of EIPA provides the legal and institutional leverage needed to open '
        'domain and infrastructure negotiations with EthioTelecom.',
        size=10, color=TEXT_MID, space_after=6)
    add_table(doc,
        headers=['Milestone', 'Target', 'Notes'],
        rows=[
            ['Complete EIPA copyright registration', 'Jun 2026  (full month)', 'Filing in progress ✅ — top priority'],
            ['Submit INSA certification', 'Jun 2026', 'Code 100% ready — run concurrently with EIPA'],
            ['Negotiate onekof.et domain — EthioTelecom', 'Jul 2026', 'EIPA completion strengthens negotiating position'],
            ['Provision EthioTelecom VM', 'Jul–Aug 2026', 'After domain agreement is secured'],
            ['Deploy Tier 2 + DNS cutover to onekof.et', 'Aug 2026', 'VM + domain live'],
            ['Email setup (Resend + .et domain)', 'Aug 2026', 'After domain live'],
            ['App Store (iOS) submission', 'Jul 2026', 'Apple Dev active ✅ — EAS build + screenshots'],
            ['Play Store (Android) submission', 'Jul 2026', 'Google Play active ✅ — EAS build + screenshots'],
            ['Internal beta — pilot organizations', 'Sep 2026', 'Tier 2 stable + INSA cert in process'],
            ['External beta / public launch', 'Oct 2026', 'Domain live + INSA cert received'],
            ['First government / sovereign pilot', 'Q1 2027', 'Post-election cycle + full INSA certification'],
        ],
        col_widths=[2.8, 1.7, 2.3],
        header_bg=TEAL
    )

    # ── 7. Competitive Position ──────────────────────────────────────────────
    section_heading(doc, '7.  Competitive Position')
    add_table(doc,
        headers=['Feature', 'Onekof', 'Jira', 'Asana', 'Trello'],
        rows=[
            ['Ethiopian language support (5 langs)', '✅', '❌', '❌', '❌'],
            ['Ethiopian calendar (Ge\'ez months)', '✅', '❌', '❌', '❌'],
            ['ETB currency default', '✅', '❌', '❌', '❌'],
            ['On-premise / air-gapped deployment', '✅', 'Enterprise only', '❌', '❌'],
            ['Government data residency', '✅ (Tiers 1 & 2)', '❌', '❌', '❌'],
            ['INSA compliance (100%)', '✅', 'N/A', 'N/A', 'N/A'],
            ['Local ETB pricing', '✅', '❌', '❌', '❌'],
            ['Mobile app (iOS + Android)', '✅', '✅ (EN only)', '✅ (EN only)', '✅ (EN only)'],
        ],
        col_widths=[3.0, 1.0, 1.0, 1.0, 1.0],
        header_bg=SLATE,
        status_col=1
    )

    add_para(doc, space_after=6)
    add_highlight_box(doc,
        '  Onekof is the only enterprise project management platform built for Ethiopian '
        'compliance, language, and data sovereignty requirements.',
        bg=TEAL, fg=WHITE, size=11)

    # ── 8. Risk Register ─────────────────────────────────────────────────────
    section_heading(doc, '8.  Risk Register')
    add_table(doc,
        headers=['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        rows=[
            ['EthioTelecom VM provisioning delays', 'Medium', 'High', 'Use Massano rig as interim Tier 2 fallback'],
            ['INSA certification timeline unknown', 'Medium', 'High', 'Submit immediately — code is 100% ready'],
            ['onekof.et domain unavailable', 'Low', 'High', 'Fallback: onekof.com.et or negotiate'],
            ['No paying customers at launch', 'Medium', 'High', 'Start NGO / private sector pilot outreach now'],
            ['App Store review rejection', 'Low', 'Medium', 'Pre-review against Apple guidelines'],
            ['EIPA filing delay', 'High', 'Medium', 'Draft Co-Owner agreement ASAP — does not block go-live'],
        ],
        col_widths=[2.5, 1.1, 0.9, 2.3],
        header_bg=SLATE
    )

    # ── 9. What Founders Need to Do ──────────────────────────────────────────
    section_heading(doc, '9.  What Founders & Stakeholders Need to Do')

    section_heading(doc, 'June 2026 — Primary Focus', level=2)
    for item in [
        '1.  Drive EIPA copyright registration to completion — the institutional milestone that unlocks EthioTelecom negotiations',
        '2.  Submit INSA certification — run concurrently with EIPA; engineering is 100% ready, nothing left to code',
        '3.  Finalize Co-Owner / Co-Founder agreement (Oli Tamrat Oli ↔ DAPS Analytics) — required to close EIPA final deposit',
        '4.  Continue App Store (iOS) + Play Store (Android) submissions — both developer accounts active ✅',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(item)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = TEAL

    section_heading(doc, 'July 2026 — Infrastructure', level=2)
    for item in [
        'Open domain registration negotiation with EthioTelecom — EIPA completion provides legal standing',
        'Provision EthioTelecom VM (Ubuntu 22.04, 4 vCPU, 8 GB RAM, 100 GB SSD) — Tier 2 runbook is ready',
        'Decide Cloudflare SSL approach — $10/mo ACM or move DNS to Vercel (free wildcard)',
    ]:
        bullet(doc, item, color=TEXT_DARK)

    section_heading(doc, 'August–September 2026 — Activation', level=2)
    for item in [
        'Deploy Tier 2 on EthioTelecom VM + DNS cutover to onekof.et',
        'Identify 3–5 pilot organizations (NGOs, ministries, private companies) for internal beta',
        'Anthropic API key — enables AI document processing and receipt analysis (post-launch)',
    ]:
        bullet(doc, item, color=TEXT_DARK)

    section_heading(doc, 'Post-Election (Q1 2027) — Government Tier', level=2)
    for item in [
        'Government LOI / pilot MOU — required before Tier 1 sovereign deployment begins',
        'Scale pilot organisations → paying customers',
        'Consider patent filings for the 4 identified innovations (see ONEKOF_IP_CLAIMS_ASSESSMENT.md)',
    ]:
        bullet(doc, item, color=TEXT_DARK)

    # ── Closing Statement ────────────────────────────────────────────────────
    section_heading(doc, '10.  Summary Statement')
    add_para(doc, space_after=6)
    add_highlight_box(doc,
        '  Onekof PM is technically complete, security-hardened to 100% INSA compliance, '
        'and ready to deploy on Ethiopian infrastructure the moment the domain is registered '
        'and the VM is provisioned.  The engineering team has delivered everything required '
        'for go-live.  The path from here to launch is entirely in the hands of the business — '
        'domain, server, certification submission, and pilot outreach.',
        bg=TEAL, fg=WHITE, size=11)

    add_para(doc, space_after=20)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_border_bottom(footer, color='1C8C7D', size=6)
    fr = footer.add_run('Onekof PM  ·  DAPS Analytics  ·  Confidential  ·  May 23, 2026')
    fr.font.name  = 'Calibri'
    fr.font.size  = Pt(8)
    fr.font.color.rgb = TEXT_MID

    path = r'C:\Users\olita\onekof-platform\ONEKOF_PROJECT_READINESS_REPORT.docx'
    doc.save(path)
    print(f'Saved: {path}')
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 — MIGRATION & DEPLOYMENT PLAN
# ═══════════════════════════════════════════════════════════════════════════════

def build_deployment_plan():
    doc = Document()
    set_page_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    add_cover_page(
        doc,
        title='Migration & Deployment Plan',
        subtitle='Technical Reference — Wave 5 Complete Edition',
        date_str='May 23, 2026',
        classification='Internal / Confidential'
    )

    # ── Status Banner ────────────────────────────────────────────────────────
    add_highlight_box(doc,
        '  Status:  Wave 5 INSA Hardening Complete  ·  Tier 3 Production-Grade  ·  Tier 2 Ready to Deploy',
        bg=TEAL, fg=WHITE, size=11)
    add_para(doc, space_after=8)

    # ── 1. Current State ─────────────────────────────────────────────────────
    section_heading(doc, '1.  Current State Inventory')
    add_table(doc,
        headers=['Component', 'Setup', 'Status'],
        rows=[
            ['Web App', 'Next.js 14, App Router, standalone output', '✅ Built, Docker-proven'],
            ['Mobile App', 'React Native + Expo SDK 54 (EAS)', '✅ Feature-complete'],
            ['Database', 'Prisma + PostgreSQL 15 — 14 migrations', '✅ Applied'],
            ['Auth', 'NextAuth.js + credentials + JWT', '✅ Working'],
            ['File Storage', 'Local-fs driver + AES-256-GCM encryption', '✅ Production-ready'],
            ['Redis', 'Upstash (cloud) / Redis 7 (self-hosted)', '✅ Ready'],
            ['Docker Image', 'Multi-stage standalone, ~408 MB', '✅ Proven locally'],
            ['CI/CD', 'GitHub Actions — CI + Deploy to Vercel', '✅ Both green'],
            ['Nginx / SSL', 'Caddyfile — auto Let\'s Encrypt', '✅ Committed'],
            ['Error Tracking', 'Sentry — onekof-web + onekof-mobile', '✅ Both live'],
            ['TypeScript build', 'ignoreBuildErrors: false, strict mode', '✅ 0 errors'],
            ['INSA compliance', 'P1–P6 all closed — Wave 5', '✅ 100%'],
            ['.et Domain', 'Not registered', '❌ Oli action required'],
        ],
        col_widths=[2.0, 3.0, 2.0],
        header_bg=SLATE,
        status_col=2
    )

    # ── 2. Applied Migrations ────────────────────────────────────────────────
    section_heading(doc, '2.  Applied Database Migrations')
    add_table(doc,
        headers=['Migration', 'Purpose'],
        rows=[
            ['0_init', 'Baseline — full schema (1,389 lines)'],
            ['20260304062200_add_automation_rules', 'Automation rules engine'],
            ['20260407_add_wiki_models', 'Wiki / Knowledge base'],
            ['20260409_add_contractor_role', 'Contractor role for RBAC'],
            ['20260409_add_task_links', 'Task dependency linking'],
            ['20260410_add_backlog_status', 'Backlog issue status'],
            ['20260410_project_visibility_default_public', 'Default project visibility'],
            ['20260410_revert_project_visibility_default', 'Revert visibility default'],
            ['20260411120000_portability_wave1', 'Three-tier architecture foundation'],
            ['20260412_add_admin_audit_log', 'Admin action audit log (Wave 3)'],
            ['20260412_remove_rate_limit_table', 'RateLimit → Redis (Wave 3)'],
            ['20260417_add_notifications', 'Notification read-state model (Wave 4)'],
            ['20260418_add_push_tokens', 'Mobile push notification tokens (Wave 4)'],
            ['20260505_add_org_audit_log', 'OrgAuditLog — INSA compliance (Wave 4)'],
        ],
        col_widths=[3.8, 3.0],
        header_bg=TEAL
    )

    # ── 3. Gap Analysis ──────────────────────────────────────────────────────
    section_heading(doc, '3.  Gap Analysis — All Blocking Gaps Resolved')
    add_table(doc,
        headers=['#', 'Gap', 'Fix', 'Status'],
        rows=[
            ['G1', 'No CI/CD pipeline', 'GitHub Actions workflows created', '✅ DONE — 2026-05-22'],
            ['G2', 'No production Docker Compose', 'docker-compose.prod.yml created', '✅ DONE — 2026-05-22'],
            ['G3', 'No Nginx/Caddy config', 'Caddyfile with auto-SSL created', '✅ DONE — 2026-05-22'],
            ['G4', 'ignoreBuildErrors: true', '5 TS errors fixed, strict mode on', '✅ DONE — 2026-05-23'],
            ['G5', 'No GitHub secrets', 'All secrets configured', '✅ DONE — 2026-05-22'],
            ['G6', 'No automated DB backup', 'GPG-encrypted cron backup script', '✅ DONE — 2026-05-22'],
        ],
        col_widths=[0.4, 2.3, 2.3, 1.8],
        header_bg=SLATE,
        status_col=3
    )

    # ── 4. INSA Security Status ──────────────────────────────────────────────
    section_heading(doc, '4.  INSA Security Compliance — 100%')
    add_table(doc,
        headers=['Priority', 'Requirement', 'Implementation', 'Status'],
        rows=[
            ['P1', 'CSRF origin validation', 'middleware.ts enforceCsrfOrigin()', '✅ Closed — Wave 5'],
            ['P2', 'Admin rate limiting (60/min/IP)', 'middleware.ts checkAdminRateLimit()', '✅ Closed — Wave 5'],
            ['P3', 'Audit log immutability', 'DELETE → 405 Method Not Allowed', '✅ Closed — Wave 5'],
            ['P4', 'AES-256-GCM blob encryption', 'local-fs.ts + BLOB_ENCRYPTION_KEY in Vercel', '✅ Closed — Wave 5'],
            ['P5', 'Session invalidation on password change', 'invalidateAllUserSessions() in change-password', '✅ Closed — Wave 5'],
            ['P6', 'HTTP security headers', 'CSP, X-Frame, HSTS, Referrer-Policy in middleware', '✅ Closed — Wave 2'],
            ['T1-7', 'Offline USB installer package', 'Build after Tier 2 live', '⏳ Pending'],
            ['T1-8', 'Amharic ops runbook', 'English runbook done — translation needed', '⏳ Pending'],
            ['T1-9', 'Compliance / legal docs', 'Data residency letter for procurement', '❌ Not started'],
        ],
        col_widths=[0.8, 2.0, 2.5, 1.5],
        header_bg=TEAL,
        status_col=3
    )

    # ── 5. Vercel Environment Variables ─────────────────────────────────────
    section_heading(doc, '5.  Vercel Production Environment Variables')
    add_table(doc,
        headers=['Variable', 'Purpose', 'Status'],
        rows=[
            ['DATABASE_URL', 'Supabase pooler (pgbouncer=true&connection_limit=1)', '✅ Set'],
            ['DIRECT_URL', 'Supabase direct (for migrations)', '✅ Set'],
            ['NEXTAUTH_SECRET', '32-byte random secret for JWT signing', '✅ Set'],
            ['NEXTAUTH_URL', 'https://onekof.com', '✅ Set'],
            ['REDIS_URL / UPSTASH_*', 'Upstash Redis for rate limiting', '✅ Set'],
            ['BLOB_READ_WRITE_TOKEN', 'Vercel Blob storage token', '✅ Set'],
            ['BLOB_ENCRYPTION_KEY', '32-byte AES key for file encryption', '✅ Set — 2026-05-23'],
            ['RESEND_API_KEY', 'Email delivery (Resend)', '✅ Set'],
            ['SENTRY_DSN', 'Server-side error tracking (onekof-web)', '✅ Set — 2026-05-23'],
            ['NEXT_PUBLIC_SENTRY_DSN', 'Client/edge error tracking (onekof-web)', '✅ Set — 2026-05-23'],
            ['ADMIN_SECRET', 'Admin panel authentication', '✅ Set'],
            ['ADMIN_USERS', 'Admin bcrypt-hashed credentials', '✅ Set'],
        ],
        col_widths=[2.5, 2.8, 1.5],
        header_bg=SLATE,
        status_col=2
    )

    # ── 6. Open Items Tracker ────────────────────────────────────────────────
    section_heading(doc, '6.  Open Items Tracker')
    add_table(doc,
        headers=['#', 'Item', 'Owner', 'Status'],
        rows=[
            ['OI-1', 'Negotiate onekof.et domain — EthioTelecom (after EIPA)', 'DAPS Analytics', '⏳ Jul 2026 — after EIPA'],
            ['OI-2', 'GitHub Actions CI workflow', 'Oli', '✅ DONE — 2026-05-22'],
            ['OI-3', 'docker-compose.prod.yml', 'Oli', '✅ DONE — 2026-05-22'],
            ['OI-4', 'Caddyfile for Tier 2 SSL', 'Oli', '✅ DONE — 2026-05-22'],
            ['OI-5', 'GitHub repository secrets', 'Oli', '✅ DONE — 2026-05-22'],
            ['OI-6', 'Wave 5: Close P1–P5 INSA gaps', 'Oli', '✅ DONE — 2026-05-23'],
            ['OI-7', 'Tier 1 offline USB delivery package', 'Oli', '⏳ Pending'],
            ['OI-8', 'Fix TypeScript errors (ignoreBuildErrors)', 'Oli', '✅ DONE — 2026-05-23'],
            ['OI-9', 'Load test (k6) before beta', 'Oli', '⏳ Pending'],
            ['OI-10', 'Resend domain for onekof.et', 'Oli', '⏳ After domain live'],
            ['OI-11', 'Enable AI features (Anthropic key)', 'Oli', '⏳ Post-launch'],
            ['OI-12', 'Submit INSA certification', 'DAPS Analytics', '⏳ Jun 2026 — code ready'],
            ['OI-13', 'Amharic ops runbook translation', 'Oli', '⏳ Pending'],
            ['OI-14', 'SENTRY_DSN + NEXT_PUBLIC_SENTRY_DSN to Vercel', 'Oli', '✅ DONE — 2026-05-23'],
            ['OI-15', 'BLOB_ENCRYPTION_KEY to Vercel', 'Oli', '✅ DONE — 2026-05-23'],
            ['OI-16', 'Custom 404 not-found page', 'Oli', '✅ DONE — 2026-05-23'],
        ],
        col_widths=[0.5, 3.2, 0.8, 2.3],
        header_bg=SLATE,
        status_col=3
    )

    # ── 7. Deployment Sequence ────────────────────────────────────────────
    section_heading(doc, '7.  Recommended Deployment Sequence')
    add_table(doc,
        headers=['Phase', 'Action', 'Target'],
        rows=[
            ['✅ DONE', 'CI/CD, GitHub secrets, docker-compose.prod.yml, Caddyfile', '2026-05-22'],
            ['✅ DONE', 'TypeScript strict, Wave 5 INSA P1–P5, BLOB key, Sentry, 404 page', '2026-05-23'],
            ['IN PROGRESS', 'Complete EIPA registration + Submit INSA certification concurrently', 'Jun 2026'],
            ['⏳ PLANNED', 'Negotiate onekof.et domain with EthioTelecom + Provision VM', 'Jul 2026'],
            ['⏳ PLANNED', 'Deploy Tier 2 → DNS cutover + Resend email setup + smoke tests', 'Aug 2026'],
            ['⏳ PLANNED', 'Load testing + Internal beta (pilot orgs) + App Store live', 'Sep 2026'],
            ['⏳ PLANNED', 'External beta / public launch + Tier 1 USB package build', 'Oct 2026'],
            ['⏳ PLANNED', 'First government/sovereign pilot — post-election cycle', 'Q1 2027'],
        ],
        col_widths=[1.4, 4.2, 1.2],
        header_bg=SLATE,
        status_col=0
    )
    add_para(doc, space_after=4)


    # ── 8. Security Checklist ────────────────────────────────────────────────
    section_heading(doc, '8.  Security Checklist')
    add_table(doc,
        headers=['Check', 'Status', 'Notes'],
        rows=[
            ['No secrets in git history', '✅ Verified', 'git log --all clean'],
            ['NEXTAUTH_SECRET ≥ 32 bytes', '✅ Verified', 'Confirmed 32 bytes'],
            ['DB ports not exposed to internet', '✅ Verified', 'expose: (inter-container only)'],
            ['Rate limiting on login', '✅ Active', 'Account lockout enabled'],
            ['CSRF on all mutation APIs', '✅ Active', 'Wave 5 — enforceCsrfOrigin()'],
            ['Session invalidation on pw change', '✅ Active', 'Wave 5 — invalidateAllUserSessions()'],
            ['Audit log append-only', '✅ Active', 'Wave 5 — DELETE → 405'],
            ['AES-256-GCM blob encryption', '✅ Active', 'Wave 5 — BLOB_ENCRYPTION_KEY in Vercel'],
            ['Security headers (CSP, HSTS etc.)', '✅ Active', 'middleware.ts — all headers set'],
            ['All 141 API routes auth-guarded', '✅ Verified', 'Full audit 2026-05-23'],
            ['.env files not in git', '✅ Verified', '.gitignore confirmed'],
            ['Sentry error tracking active', '✅ Active', 'Web + mobile projects live'],
            ['Postgres/Redis passwords strong', '⏳ Pending', 'Set when Tier 2 VM provisioned'],
            ['SSL Labs rating A or better', '⏳ Pending', 'When onekof.et domain live'],
            ['Backup restoration tested', '⏳ Pending', 'When Tier 2 live'],
        ],
        col_widths=[2.8, 1.3, 2.7],
        header_bg=SLATE,
        status_col=1
    )

    add_para(doc, space_after=16)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_border_bottom(footer, color='1C8C7D', size=6)
    fr = footer.add_run('Onekof PM  ·  Migration & Deployment Plan  ·  DAPS Analytics  ·  May 23, 2026')
    fr.font.name  = 'Calibri'
    fr.font.size  = Pt(8)
    fr.font.color.rgb = TEXT_MID

    path = r'C:\Users\olita\onekof-platform\ONEKOF_MIGRATION_DEPLOYMENT_PLAN.docx'
    doc.save(path)
    print(f'Saved: {path}')
    return path


# ─── Run ─────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    p1 = build_readiness_report()
    p2 = build_deployment_plan()
    print('\nBoth documents generated successfully.')
    print(f'  1. {p1}')
    print(f'  2. {p2}')
