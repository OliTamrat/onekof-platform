"""
Generates MIGRATION_DEPLOYMENT_PLAN.docx from the migration plan content.
Run: python scripts/generate_migration_plan_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND_BLUE   = RGBColor(0x1A, 0x56, 0xDB)
BRAND_DARK   = RGBColor(0x11, 0x18, 0x27)
ACCENT_GREEN = RGBColor(0x05, 0x96, 0x69)
ACCENT_AMBER = RGBColor(0xD9, 0x77, 0x06)
TABLE_HEADER = RGBColor(0x1E, 0x40, 0xAF)
TABLE_ALT    = RGBColor(0xEF, 0xF6, 0xFF)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','bottom','left','right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:color'), 'BFDBFE')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_BLUE
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:color'), '1A56DB')
    pBdr.append(bot)
    pPr.append(pBdr)


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.color.rgb = BRAND_DARK


def h3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_BLUE


def body(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def bullet(doc, text, done=False, pending=False):
    para  = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(2)
    color = ACCENT_GREEN if done else (ACCENT_AMBER if pending else RGBColor(0x1F,0x29,0x37))
    run = para.add_run(('✓  ' if done else '☐  ') + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = color


def code_block(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.left_indent = Inches(0.3)
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F1F5F9')
        para._p.get_or_add_pPr().append(shd)


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        set_cell_bg(hdr_cells[i], TABLE_HEADER)
        set_cell_border(hdr_cells[i])
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], TABLE_ALT)
            set_cell_border(cells[c_idx])
            run = cells[c_idx].paragraphs[0].add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            v = str(val)
            if 'Done' in v or 'DONE' in v:
                run.font.color.rgb = ACCENT_GREEN
                run.bold = True
            elif 'Pending' in v or 'Blocked' in v or 'Open' in v:
                run.font.color.rgb = ACCENT_AMBER

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ── Cover ──────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(24)
r = title_para.add_run('Onekof PM')
r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(32)
r.font.color.rgb = BRAND_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run('Migration & Deployment Plan')
r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(20)
r.font.color.rgb = BRAND_DARK

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(meta, 'Author: Oli T. Oli / DAPS Analytics     |     2026-05-23     |     Internal / Confidential',
        size=10, color=RGBColor(0x6B, 0x72, 0x80))

status_p = doc.add_paragraph()
status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
status_p.paragraph_format.space_before = Pt(6)
add_run(status_p, 'Status: Tier 3 CI/CD Complete — Tier 2 Ready to Deploy',
        bold=True, size=11, color=ACCENT_GREEN)

doc.add_paragraph()

# ── 1. Current State ───────────────────────────────────────────────────────
h1(doc, '1. Current State Inventory')

make_table(doc,
    headers=['Component', 'Current Setup', 'Status'],
    rows=[
        ['Web App',       'Next.js 14, App Router, standalone output',   'Built, Docker-proven'],
        ['Mobile App',    'React Native + Expo (EAS)',                    'Feature-complete'],
        ['Database',      'Prisma 5 + PostgreSQL 15',                    '14 migrations applied'],
        ['Auth',          'NextAuth.js v5 + credentials',                'Working'],
        ['File Storage',  'local-fs driver (Wave 1)',                    'Production-ready'],
        ['Redis',         'Upstash (cloud) / Redis 7 (self-hosted)',     'Ready'],
        ['Docker Image',  'Multi-stage, ~300 MB standalone',             'Proven locally'],
        ['CI/CD',         'GitHub Actions (CI + Deploy to Vercel)',      'Done ✓'],
        ['SSL / Proxy',   'Caddyfile (auto Let\'s Encrypt)',             'Done ✓'],
        ['.et Domain',    'Not registered',                              'Pending'],
        ['INSA Security', 'P1–P6 all closed (Wave 5)',                  'Done ✓ 2026-05-23'],
    ],
    col_widths=[1.5, 3.0, 2.2]
)

h2(doc, 'Monorepo Structure')
code_block(doc, [
    'onekof-platform/',
    '  apps/',
    '    web/          # Next.js 14 — Dockerfile lives here',
    '    mobile/       # Expo React Native',
    '  packages/',
    '    database/     # Prisma schema + migrations',
    '    config/       # Shared TS / ESLint / Tailwind configs',
    '  docker-compose.prod.yml',
    '  Caddyfile',
    '  .env.production.example',
    '  scripts/setup-tier2-server.sh',
])

h2(doc, 'Applied Database Migrations (14 total)')
code_block(doc, [
    '0_init',
    '20260304062200_add_automation_rules',
    '20260407_add_wiki_models',
    '20260409_add_contractor_role',
    '20260409_add_task_links',
    '20260410_add_backlog_status',
    '20260410_project_visibility_default_public',
    '20260410_revert_project_visibility_default',
    '20260411120000_portability_wave1',
    '20260412_add_admin_audit_log',
    '20260412_remove_rate_limit_table',
    '20260417_add_notifications',
    '20260418_add_push_tokens',
    '20260505_add_org_audit_log',
])

# ── 2. Three-Tier Architecture ─────────────────────────────────────────────
h1(doc, '2. Three-Tier Deployment Architecture')

body(doc, 'The same Docker image runs in all three tiers. Runtime behaviour is controlled entirely by environment variables.')

make_table(doc,
    headers=['Tier', 'Environment', 'Target Date', 'Use Case'],
    rows=[
        ['Tier 3', 'Vercel + Supabase cloud',                  'LIVE',      'Global SaaS, diaspora, international'],
        ['Tier 2', 'Self-hosted EthioTelecom VM (onekof.et)', '2026-06-14', 'Ethiopian data-residency deployments'],
        ['Tier 1', 'Air-gapped sovereign (USB / ACR)',         '2026-08-01', 'Government, INSA-classified environments'],
    ],
    col_widths=[0.7, 2.5, 1.3, 2.2]
)

# ── 3. Gap Analysis ────────────────────────────────────────────────────────
h1(doc, '3. Gap Analysis')

h2(doc, 'Blocking Gaps')
make_table(doc,
    headers=['#', 'Gap', 'Impact', 'Status'],
    rows=[
        ['G1', 'No CI/CD pipeline',                            'Manual deploys',          'Done ✓'],
        ['G2', 'No production Docker Compose',                 'No Tier 2 config',        'Done ✓'],
        ['G3', 'No Caddy/SSL config',                          'No SSL for self-hosted',  'Done ✓'],
        ['G4', 'ignoreBuildErrors: true in next.config',       'TS errors silently hidden','Open'],
        ['G5', 'No GitHub secrets configured',                 'CI cannot deploy',        'Done ✓'],
        ['G6', 'No automated DB backup',                       'Data loss risk',          'Done ✓'],
    ],
    col_widths=[0.4, 2.5, 2.0, 1.8]
)

h2(doc, 'Non-Blocking Gaps')
make_table(doc,
    headers=['#', 'Gap', 'Impact', 'Status'],
    rows=[
        ['G7',  'No .et domain',                 'Non-Ethiopian URL',           'Pending'],
        ['G8',  'Email not verified end-to-end', 'Users may miss emails',       'Pending (needs domain)'],
        ['G9',  'AI features disabled',          'Feature gap vs. spec',        'Pending (post-launch)'],
        ['G10', 'No load test run',              'Unknown capacity',            'Beta week'],
        ['G11', 'TypeScript ignoreBuildErrors',  'TS errors bypass CI',         'Beta week'],
    ],
    col_widths=[0.4, 2.3, 2.0, 2.0]
)

# ── 4. Phase 1 — Tier 3 Hardening ─────────────────────────────────────────
h1(doc, '4. Phase 1 — Cloud Production Hardening (Tier 3) — COMPLETE')

body(doc, 'All steps below are done. Tier 3 is in a production-grade state.')

done_items = [
    'CI workflow: build + type check on every push to master/develop',
    'Deploy workflow: auto-deploys to Vercel on master push via vercel pull + vercel build --prod + vercel deploy --prebuilt',
    'GitHub secrets set: VERCEL_TOKEN, VERCEL_ORG_ID, VERCEL_PROJECT_ID, DIRECT_URL',
    'DATABASE_URL uses Session Pooler (port 5432) — IPv4 compatible, works from GitHub Actions',
    'DIRECT_URL uses Session Pooler — supports DDL migrations (pgbouncer blocks DDL)',
    'schema.prisma has directUrl set — Prisma migrations work correctly',
    '5 stuck migrations resolved via prisma migrate resolve --applied',
    'Wave 5 INSA security hardening: P1 CSRF, P2 admin rate limit, P3 audit log append-only, P4 blob encryption, P5 session invalidation',
]
for item in done_items:
    bullet(doc, item, done=True)

h2(doc, 'Vercel Environment Variables Required')
make_table(doc,
    headers=['Variable', 'Source', 'Notes'],
    rows=[
        ['DATABASE_URL',      'Supabase',  'Session Pooler port 6543, pgbouncer=true&connection_limit=1'],
        ['DIRECT_URL',        'Supabase',  'Session Pooler port 5432, no pgbouncer params'],
        ['NEXTAUTH_SECRET',   'Generated', 'openssl rand -base64 32'],
        ['NEXTAUTH_URL',      'Manual',    'https://onekof.com or current Vercel URL'],
        ['RESEND_API_KEY',    'Resend',    'from resend.com'],
        ['SENTRY_DSN',        'Sentry',    'from sentry.io project'],
        ['FIELD_ENCRYPTION_KEY', 'Generated', '64 hex chars — PII field encryption (audit log IPs)'],
        ['BLOB_ENCRYPTION_KEY',  'Generated', '64 hex chars — file at-rest encryption (Tier 2/1)'],
    ],
    col_widths=[2.0, 1.2, 3.5]
)

# ── 5. Phase 2 — Tier 2 ───────────────────────────────────────────────────
h1(doc, '5. Phase 2 — Self-Hosted Ethiopia (Tier 2) — Target: 2026-06-14')

h2(doc, 'Infrastructure Requirements')
make_table(doc,
    headers=['Component', 'Minimum', 'Recommended'],
    rows=[
        ['CPU',     '2 vCPU',        '4 vCPU'],
        ['RAM',     '4 GB',          '8 GB'],
        ['Disk',    '40 GB SSD',     '100 GB SSD'],
        ['OS',      'Ubuntu 22.04',  'Ubuntu 22.04 LTS'],
        ['Network', '10 Mbps',       '100 Mbps'],
        ['Ports',   '80, 443',       '80, 443, 22 (restricted to your IP)'],
    ],
    col_widths=[1.5, 2.0, 3.2]
)

h2(doc, 'Deployment Steps')
steps = [
    ('Step 2.1', 'Server provisioning — SSH in, configure UFW firewall (80, 443 open; 5432, 6379 blocked)'),
    ('Step 2.2', 'Clone repo to /opt/onekof, run: bash scripts/setup-tier2-server.sh'),
    ('Step 2.3', 'Edit .env.production — fill in NEXTAUTH_URL, RESEND_API_KEY, SENTRY_DSN'),
    ('Step 2.4', 'Update Caddyfile with real domain (onekof.et), restart Caddy for auto-SSL'),
    ('Step 2.5', 'Point DNS A record to server IP, verify Let\'s Encrypt cert'),
    ('Step 2.6', 'Run full smoke test checklist (see below)'),
    ('Step 2.7', 'Verify blob encryption: uploaded files unreadable as raw bytes on disk'),
    ('Step 2.8', 'Confirm backup cron installed: crontab -l'),
]
for step, desc in steps:
    h3(doc, step)
    body(doc, desc)

h2(doc, 'Tier 2 Smoke Test Checklist')
smoke_items = [
    'docker compose -f docker-compose.prod.yml ps — all services healthy',
    'curl -f http://localhost:3000/api/health — {"status":"ok"}',
    'https://onekof.et loads — Caddy SSL working',
    'Let\'s Encrypt certificate valid (check at ssllabs.com)',
    'User can sign up — receives verification email',
    'User can log in',
    'Project creation works',
    'Issue creation works',
    'File upload works (STORAGE_DRIVER=local-fs)',
    'File download works (decrypted transparently)',
    'Audit log records actions',
    'Rate limiting triggers after 4+ failed logins',
    'Backup cron installed (crontab -l shows onekof-daily-backup)',
    'Sentry receiving events',
]
for item in smoke_items:
    bullet(doc, item)

# ── 6. Phase 3 — Tier 1 ───────────────────────────────────────────────────
h1(doc, '6. Phase 3 — Sovereign / Air-Gapped Deployment (Tier 1) — Target: 2026-08-01')

body(doc, 'INSA compliance is now at 100% (all P1–P6 gaps closed in Wave 5, 2026-05-23). This unblocks formal INSA certification submission and government procurement conversations.')

h2(doc, 'INSA Compliance Status')
make_table(doc,
    headers=['Gap', 'Description', 'Status'],
    rows=[
        ['P1', 'CSRF Origin validation for all API mutations',           'Done ✓ — middleware.ts'],
        ['P2', 'Rate limiting on all /api/admin/* routes',              'Done ✓ — middleware.ts'],
        ['P3', 'Audit log append-only (explicit HTTP 405 on DELETE)',   'Done ✓ — audit-log/route.ts'],
        ['P4', 'AES-256-GCM at-rest encryption for blob storage',      'Done ✓ — local-fs.ts'],
        ['P5', 'Session invalidation on settings password change',      'Done ✓ — change-password/route.ts'],
        ['P6', 'HTTP security headers on self-hosted (Caddy)',          'Done ✓ — Caddyfile'],
    ],
    col_widths=[0.4, 3.5, 2.8]
)

h2(doc, 'Tier 1 USB Delivery Package')
code_block(doc, [
    'USB drive (16 GB+):',
    '  onekof-tier1-v1.0.0.tar.gz           # Docker image (~300 MB)',
    '  onekof-tier1-v1.0.0.tar.gz.sha256    # integrity hash',
    '  docker-compose.prod.yml              # stack config',
    '  .env.production.example              # env template',
    '  Caddyfile                            # reverse proxy',
    '  TIER1_DEPLOYMENT_RUNBOOK.md          # English',
    '  TIER1_DEPLOYMENT_RUNBOOK_AM.md       # Amharic',
    '  scripts/setup-tier2-server.sh        # works for Tier 1 too',
])

h2(doc, 'Government Target Segments (Priority Order)')
gov_targets = [
    ('Priority 1', 'Ministry of Innovation & Technology (MInT)', 'Digital transformation project management'),
    ('Priority 2', 'Ethiopian Investment Commission (EIC)',       'Investment facilitation project tracking'),
    ('Priority 3', 'Commercial Bank of Ethiopia (CBE)',           'Internal IT project management'),
    ('Priority 4', 'Ethio Telecom',                              'Infrastructure rollout project tracking'),
]
make_table(doc,
    headers=['Priority', 'Organization', 'Use Case'],
    rows=gov_targets,
    col_widths=[0.8, 2.5, 3.4]
)

# ── 7. CI/CD Pipeline ─────────────────────────────────────────────────────
h1(doc, '7. CI/CD Pipeline')

h2(doc, 'Branch Strategy')
make_table(doc,
    headers=['Branch', 'Trigger', 'Result'],
    rows=[
        ['master',    'Push',              'Auto-deploy to Vercel production (Tier 3)'],
        ['develop',   'Push',              'Auto-deploy to Vercel preview'],
        ['feature/*', 'PR open/update',    'CI build only — no deploy'],
        ['v*.*.*',    'Tag push',          'Docker image built + pushed to GHCR (Tier 2/1)'],
    ],
    col_widths=[1.2, 1.5, 4.0]
)

h2(doc, 'Workflow Files')
make_table(doc,
    headers=['File', 'Purpose'],
    rows=[
        ['.github/workflows/ci.yml',                'Build + type check + unit tests on every push'],
        ['.github/workflows/deploy-production.yml', 'Vercel pull → DB migrate → vercel build → vercel deploy --prebuilt'],
        ['.github/workflows/docker-build.yml',      'Build + push Docker image to GHCR on version tags (v*.*.*)'],
    ],
    col_widths=[3.0, 3.7]
)

h2(doc, 'GitHub Secrets Required')
make_table(doc,
    headers=['Secret', 'Value'],
    rows=[
        ['VERCEL_TOKEN',      'From vercel.com/account/tokens'],
        ['VERCEL_ORG_ID',     'team_1MQSZxhdEsQsh7Z9DViqHiuc'],
        ['VERCEL_PROJECT_ID', 'prj_STTRolqEQMNvTLH3khdFfwE0eoQh'],
        ['DIRECT_URL',        'Session Pooler port 5432 — used by prisma migrate deploy in CI'],
    ],
    col_widths=[2.0, 4.7]
)

# ── 8. Data Migration ──────────────────────────────────────────────────────
h1(doc, '8. Data Migration (Supabase → Self-Hosted)')

h2(doc, 'When to Run')
body(doc, 'When a customer requests to move their data from Tier 3 (Supabase cloud) to Tier 2 (self-hosted Ethiopia). Maintenance window: approximately 2 hours.')

h2(doc, 'Export from Supabase')
code_block(doc, [
    'vercel env pull .env.supabase --environment=production',
    '',
    'pg_dump \\',
    '  --no-owner --no-privileges \\',
    '  --format=custom \\',
    '  "$(grep DIRECT_URL .env.supabase | cut -d= -f2-)" \\',
    '  -f onekof-export-$(date +%Y%m%d).dump',
])

h2(doc, 'Import to Tier 2')
code_block(doc, [
    'scp onekof-export-*.dump ubuntu@<server-ip>:/opt/onekof/',
    '',
    'docker compose -f docker-compose.prod.yml exec -T postgres \\',
    '  pg_restore --no-owner --no-privileges \\',
    '  -U onekof -d onekof \\',
    '  < /opt/onekof/onekof-export-*.dump',
])

h2(doc, 'Verify Data Integrity')
code_block(doc, [
    'docker compose exec postgres psql -U onekof -d onekof -c "',
    '  SELECT',
    '    (SELECT COUNT(*) FROM users) as users,',
    '    (SELECT COUNT(*) FROM organizations) as orgs,',
    '    (SELECT COUNT(*) FROM projects) as projects,',
    '    (SELECT COUNT(*) FROM issues) as issues;"',
])

# ── 9. Production Cutover ──────────────────────────────────────────────────
h1(doc, '9. Production Cutover Checklist')

h2(doc, 'Pre-Cutover (48 hours before)')
pre_items = [
    'Final full database backup from current environment',
    'Verify backup can be restored (test restore to staging)',
    'All environment variables documented and secured',
    'Docker image tagged with release version (e.g., v1.0.0)',
    'All pending migrations applied to production DB',
    'Smoke test on staging passes',
    'Sentry error tracking active',
    'Uptime monitor configured (UptimeRobot or similar)',
    'Rollback plan reviewed',
]
for item in pre_items:
    bullet(doc, item)

h2(doc, 'Cutover Day')
cutover_items = [
    'Put app in maintenance mode (NEXT_PUBLIC_MAINTENANCE=true)',
    'Final database backup',
    'Deploy new image / environment',
    'Run prisma migrate deploy',
    'Verify all services healthy (docker compose ps)',
    'Run smoke tests',
    'Switch DNS to new server',
    'Verify SSL certificate valid',
    'Remove maintenance mode',
    'Monitor Sentry for 2 hours',
]
for item in cutover_items:
    bullet(doc, item)

h2(doc, 'Post-Cutover (48 hours)')
post_items = [
    'Zero P0/P1 errors in Sentry',
    'Uptime monitor — all green',
    'Server resource usage normal (CPU < 70%, RAM < 80%)',
    'Automated backups ran at scheduled time',
    'Email delivery working (check Resend dashboard)',
]
for item in post_items:
    bullet(doc, item)

# ── 10. Rollback ───────────────────────────────────────────────────────────
h1(doc, '10. Rollback Plan')

h2(doc, 'Tier 3 — Vercel (Instant, zero-downtime)')
code_block(doc, ['vercel rollback'])

h2(doc, 'Tier 2 — Self-Hosted')
code_block(doc, [
    'docker compose -f docker-compose.prod.yml stop onekof-web',
    'nano .env.production  # set DOCKER_IMAGE=onekof-web:vPREVIOUS',
    'docker compose -f docker-compose.prod.yml pull onekof-web',
    'docker compose -f docker-compose.prod.yml up -d onekof-web',
])

h2(doc, 'Database Rollback')
code_block(doc, [
    'docker compose exec -T postgres psql -U onekof -d onekof \\',
    '  < /opt/backups/onekof-<date>.sql.gz',
])

# ── 11. Security Checklist ─────────────────────────────────────────────────
h1(doc, '11. Security Checklist Before Launch')

security_items = [
    ('No secrets in git history', 'git log --all --full-history -- "*.env"'),
    ('No hard-coded credentials in source code', None),
    ('NEXTAUTH_SECRET is at least 32 random bytes', None),
    ('Postgres and Redis passwords are strong (16+ chars)', None),
    ('Ports 5432 and 6379 NOT exposed to internet', 'check UFW + docker-compose'),
    ('SSH key-based auth only (disable password auth)', None),
    ('Rate limiting verified (5 failed logins → lockout)', None),
    ('Audit log capturing actions and append-only (405 on DELETE)', None),
    ('CSRF Origin validation active for all API mutations', None),
    ('Admin API rate-limited at 60 req/min', None),
    ('Blob encryption active (files unreadable as raw bytes)', None),
    ('Security headers: A rating at securityheaders.com', None),
    ('SSL Labs: A+ rating at ssllabs.com', None),
    ('Sentry DSN is set and receiving events', None),
    ('Backup restoration tested at least once', None),
    ('.env.production not committed to git (verify .gitignore)', None),
]
for item, note in security_items:
    text = item + (f'  [{note}]' if note else '')
    bullet(doc, text)

# ── 12. Open Items ─────────────────────────────────────────────────────────
h1(doc, '12. Open Items Tracker')

make_table(doc,
    headers=['#', 'Item', 'Owner', 'Deadline', 'Status'],
    rows=[
        ['OI-1',  'Register onekof.et domain',                  'Oli', 'Jun 7',  'Pending — do first'],
        ['OI-2',  'GitHub Actions CI workflow',                  'Oli', 'Done',   'Done'],
        ['OI-3',  'docker-compose.prod.yml',                     'Oli', 'Done',   'Done'],
        ['OI-4',  'Caddyfile for Tier 2',                        'Oli', 'Done',   'Done'],
        ['OI-5',  'GitHub repository secrets',                   'Oli', 'Done',   'Done'],
        ['OI-6',  'Wave 5 — close P1–P5 INSA gaps',             'Oli', 'Done',   'Done'],
        ['OI-7',  'Provision EthioTelecom VM',                   'Oli', 'Jun 7',  'Pending — HIGH'],
        ['OI-8',  'Fix TypeScript errors (ignoreBuildErrors)',   'Oli', 'Jun 21', 'Pending'],
        ['OI-9',  'Load test (k6) before beta',                  'Oli', 'Jun 21', 'Pending'],
        ['OI-10', 'Provision Resend domain for onekof.et',       'Oli', 'Jun 14', 'Blocked on OI-1'],
        ['OI-11', 'Submit INSA certification',                   'Oli', 'May 31', 'Pending — HIGH'],
        ['OI-12', 'Build Tier 1 offline USB delivery package',   'Oli', 'Jul 20', 'Pending'],
        ['OI-13', 'Amharic ops runbook',                         'Oli', 'Jul 20', 'Pending'],
    ],
    col_widths=[0.5, 2.7, 0.6, 0.8, 1.6]
)

# ── 13. Deployment Sequence ────────────────────────────────────────────────
h1(doc, '13. Recommended Deployment Sequence')

make_table(doc,
    headers=['Timeline', 'Action'],
    rows=[
        ['Done ✓ (2026-05-23)', 'CI/CD, GitHub secrets, docker-compose.prod.yml, Caddyfile, Wave 5 INSA hardening'],
        ['Week 1 (May 25–31)',  'Register onekof.et — Provision EthioTelecom VM — Submit INSA certification'],
        ['Week 2 (Jun 1–7)',    'Server setup: run setup-tier2-server.sh, configure .env.production, smoke test'],
        ['Week 3 (Jun 8–14)',   'DNS cutover → onekof.et — Caddy SSL — Resend email verification — Tier 2 Go-Live'],
        ['Week 4 (Jun 15–21)', 'Internal beta — k6 load test — TypeScript cleanup — monitoring setup'],
        ['Week 5 (Jun 22–28)', 'INSA follow-up — public beta prep — mobile APK distribution'],
        ['Week 6 (Jun 29–Jul 5)', 'Public beta launch — LinkedIn + Telegram announcement'],
        ['Week 7–8 (Jul 6–19)', 'Government pilot outreach (MInT, EIC, CBE, Ethio Telecom)'],
        ['Week 9–12 (Jul 20–Aug 15)', 'Tier 1 USB package — first sovereign government deployment'],
    ],
    col_widths=[2.2, 4.5]
)

# ── Footer ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer,
        'Document Owner: Oli T. Oli / DAPS Analytics   |   Next Review: 2026-06-14   |   Source of Truth: MIGRATION_DEPLOYMENT_PLAN.md',
        size=9, color=RGBColor(0x9C, 0xA3, 0xAF))

# ── Save ───────────────────────────────────────────────────────────────────
output_path = 'MIGRATION_DEPLOYMENT_PLAN.docx'
doc.save(output_path)
print(f'Done. Saved to: {output_path}')
