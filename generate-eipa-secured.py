"""
EIPA Secured Deposit Generator
==============================

Generates a SECOND version of the deposit with additional security measures
layered on top of the standard `DEPOSIT/` folder. This secured version is
intended to be offered to EIPA as an alternative if they accept it — or held
in reserve if the standard deposit is preferred.

Security measures applied:
  (a) Password-protected (AES-256) ZIP of the full source archive
  (b) SHA-256 manifest of every file in the secured deposit (tamper detection)
  (c) Reduced "minimal" source deposit — first ~25 / last ~25 pages only
  (d) Filing Agent NDA (Word document) for whoever files on the author's behalf
  (e) README explaining every artifact in the secured deposit

Output: DEPOSIT_SECURED/ folder at the repo root.
  - 00_FORMS/                             (the 4 filing forms, copied)
  -     Filing_Agent_NDA.docx              (new NDA for the filing agent)
  - 01_MINIMAL_SOURCE/                    (reduced deposit)
  -     First_25_Pages.docx
  -     Last_25_Pages.docx
  -     Minimal_Source_README.txt
  - onekof-platform-source.zip.aes        (password-protected full ZIP)
  - SHA256_MANIFEST.txt                   (plain text hash list)
  - SHA256_MANIFEST.docx                  (formatted hash list)
  - ZIP_PASSWORD.txt                      (NOT to be placed on the CD!)
  - README.txt                            (index explaining every file)

Author: Oli Tamrat Oli <oli.oli@udc.edu>
Date:   2026-04-11
"""

import hashlib
import os
import secrets
import shutil
import string
import subprocess
from datetime import datetime
from pathlib import Path

import pyzipper
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO_ROOT = Path(__file__).parent
DEPOSIT_DIR = REPO_ROOT / "DEPOSIT"
SECURED_DIR = REPO_ROOT / "DEPOSIT_SECURED"
FORMS_DIR = SECURED_DIR / "00_FORMS"
MINIMAL_DIR = SECURED_DIR / "01_MINIMAL_SOURCE"

TEAL = RGBColor(0x1C, 0x8C, 0x7D)
DARK = RGBColor(0x22, 0x27, 0x2B)
GREY = RGBColor(0x77, 0x77, 0x77)

TODAY = datetime.now().strftime("%B %d, %Y")
TODAY_ISO = datetime.now().strftime("%Y-%m-%d")

# Approx pages target for minimal deposit. 50 lines per printed page at 8pt
# Consolas is the historical Onekof rendering cadence.
LINES_PER_PAGE = 50
TARGET_PAGES_PER_SIDE = 25
TARGET_LINES_PER_SIDE = LINES_PER_PAGE * TARGET_PAGES_PER_SIDE  # 1250

# ---------------------------------------------------------------------------
# Shared docx styling
# ---------------------------------------------------------------------------

def set_document_defaults(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = TEAL
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1C8C7D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL


def add_paragraph(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if italic:
        run.italic = True
    if bold:
        run.bold = True


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].text = item
            p.runs[0].font.size = Pt(11)
        else:
            run = p.add_run(item)
            run.font.size = Pt(11)


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_field_table(doc, rows, label_width_inches=2.3):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.style = 'Light List Accent 1'
    for i, (label, value) in enumerate(rows):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        label_cell.width = Inches(label_width_inches)
        value_cell.width = Inches(6.5 - label_width_inches)
        _shade_cell(label_cell, 'E8F5F3')

        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lrun = lp.add_run(label)
        lrun.bold = True
        lrun.font.size = Pt(10)
        lrun.font.color.rgb = DARK

        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        display = value if value else "__________________________________________"
        vrun = vp.add_run(display)
        vrun.font.size = Pt(10)
        if not value:
            vrun.font.color.rgb = GREY

        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def add_code_block(doc, content, font_size=8):
    """Render a block of source code with line numbers, one run with <w:br>."""
    lines = content.splitlines()
    if not lines:
        lines = ['(empty)']
    line_num_width = max(len(str(len(lines))), 3)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)

    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.append(rFonts)

    for i, line in enumerate(lines):
        if i > 0:
            br = OxmlElement('w:br')
            run._element.append(br)
        line_num = f"{i + 1:>{line_num_width}} | "
        display = line_num + line.replace('\t', '    ')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = _sanitize_for_xml(display)
        run._element.append(t)


def add_file_header(doc, file_path, line_count):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"FILE: {file_path}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    meta_run = p2.add_run(f"({line_count} lines)")
    meta_run.italic = True
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = GREY


def _sanitize_for_xml(text):
    if not text:
        return text
    return ''.join(
        ch for ch in text
        if ch in ('\t', '\n', '\r')
        or 0x20 <= ord(ch) <= 0xD7FF
        or 0xE000 <= ord(ch) <= 0xFFFD
    )


def add_footer(doc, label):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{label}  |  Onekof Platform — SECURED DEPOSIT  |  "
        f"Oli Tamrat Oli  |  {TODAY}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ---------------------------------------------------------------------------
# Step 1 — Copy filing forms
# ---------------------------------------------------------------------------

def copy_existing_forms():
    src = DEPOSIT_DIR / "00_FORMS"
    if not src.exists():
        raise RuntimeError(
            f"Expected forms at {src}. Run generate-eipa-forms.py first."
        )
    FORMS_DIR.mkdir(parents=True, exist_ok=True)
    copied = 0
    for f in src.iterdir():
        if f.is_file():
            shutil.copy2(f, FORMS_DIR / f.name)
            copied += 1
    return copied


# ---------------------------------------------------------------------------
# Step 2 — Generate Filing Agent NDA
# ---------------------------------------------------------------------------

def generate_filing_agent_nda(output_path):
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "Filing Agent NDA")

    add_title(
        doc,
        "NON-DISCLOSURE AGREEMENT",
        "Filing Agent — EIPA Copyright Deposit (Onekof Platform)",
    )

    add_paragraph(
        doc,
        f"This Non-Disclosure Agreement (the \u201cAgreement\u201d) is entered into "
        f"on this ______ day of ______________, 2026, by and between:",
    )

    add_field_table(doc, [
        ("Disclosing Party",
            "Oli Tamrat Oli, Ethiopian national, founder and sole author "
            "of the Onekof Platform, residing at 14823 Crossvalley Rd, "
            "Burtonsville, Maryland 20866, United States of America "
            "(email: oli.oli@udc.edu)"),
        ("Receiving Party",
            ""),
        ("Receiving Party Address",
            ""),
        ("Receiving Party ID / Passport",
            ""),
        ("Purpose of Disclosure",
            "Submission of the Onekof Platform software copyright "
            "registration application to the Ethiopian Intellectual "
            "Property Authority (EIPA) on behalf of the Disclosing Party, "
            "who resides abroad and cannot file in person."),
    ], label_width_inches=2.2)

    add_section_heading(doc, "1. Confidential Information")
    add_paragraph(
        doc,
        "\u201cConfidential Information\u201d means all information, materials, and "
        "media provided by the Disclosing Party to the Receiving Party for "
        "the purpose of the EIPA filing, including but not limited to:",
    )
    add_bullets(doc, [
        "The complete source code of the Onekof Platform "
        "(132,173 lines across 689 files)",
        "The source code deposit CD-ROM and any copies thereof",
        "The accompanying Microsoft Word filing documents "
        "(Form C-1, Software Description, Publication Date, "
        "Source Code Deposit Cover)",
        "Any encryption passwords, access keys, or authentication "
        "material provided by the Disclosing Party",
        "Business information, architecture, database schemas, "
        "trade secrets, and commercial plans disclosed in the "
        "course of the filing",
    ])

    add_section_heading(doc, "2. Obligations of the Receiving Party")
    add_paragraph(
        doc,
        "The Receiving Party shall:",
    )
    add_bullets(doc, [
        "Use the Confidential Information solely for the purpose of "
        "submitting the EIPA copyright application, and for no other "
        "purpose whatsoever.",
        "Not copy, reproduce, photograph, scan, transmit, upload, or "
        "otherwise duplicate any part of the Confidential Information "
        "except as strictly required by EIPA as part of the filing.",
        "Not disclose, share, show, or reveal the Confidential "
        "Information to any third party other than authorized EIPA "
        "officials handling the filing.",
        "Return the original CD-ROM, all Word documents, and any "
        "materials to the Disclosing Party immediately upon completion "
        "of the filing, along with the EIPA-stamped receipt.",
        "Destroy any working copies, notes, or derivative materials in "
        "the Receiving Party's possession after the filing is complete.",
        "Notify the Disclosing Party within 24 hours of any suspected "
        "unauthorized access, loss, theft, or disclosure of the "
        "Confidential Information.",
    ])

    add_section_heading(doc, "3. Exclusions")
    add_paragraph(
        doc,
        "Confidential Information does not include information that: "
        "(a) is already in the public domain at the time of disclosure "
        "through no fault of the Receiving Party; (b) was lawfully known "
        "to the Receiving Party prior to disclosure; or (c) is required "
        "to be disclosed by a court order or by Ethiopian law, provided "
        "the Receiving Party gives prompt written notice to the Disclosing "
        "Party before such disclosure.",
    )

    add_section_heading(doc, "4. Term")
    add_paragraph(
        doc,
        "The obligations under this Agreement shall commence upon the date "
        "of signing and shall remain in force indefinitely with respect to "
        "the source code of the Onekof Platform, which constitutes a trade "
        "secret of the Disclosing Party. The confidentiality obligations "
        "survive the completion of the EIPA filing.",
    )

    add_section_heading(doc, "5. Remedies")
    add_paragraph(
        doc,
        "The Receiving Party acknowledges that any breach of this Agreement "
        "may cause irreparable harm to the Disclosing Party for which "
        "monetary damages would be inadequate. Accordingly, the Disclosing "
        "Party shall be entitled to seek injunctive relief and specific "
        "performance in addition to any other remedies available under "
        "Ethiopian law, including but not limited to damages, accounting "
        "of profits, and legal costs.",
    )

    add_section_heading(doc, "6. Governing Law")
    add_paragraph(
        doc,
        "This Agreement shall be governed by and construed in accordance "
        "with the laws of the Federal Democratic Republic of Ethiopia. Any "
        "disputes arising under this Agreement shall be resolved in the "
        "courts of Addis Ababa.",
    )

    add_section_heading(doc, "7. Entire Agreement")
    add_paragraph(
        doc,
        "This Agreement constitutes the entire understanding between the "
        "parties with respect to the subject matter and supersedes all prior "
        "negotiations, representations, or agreements, whether written or "
        "oral.",
    )

    add_section_heading(doc, "8. Signatures")
    add_paragraph(
        doc,
        "By signing below, each party acknowledges that they have read, "
        "understood, and agree to be bound by the terms of this Agreement.",
    )

    add_field_table(doc, [
        ("Disclosing Party", "Oli Tamrat Oli"),
        ("Signature", ""),
        ("Date", ""),
        ("Place", "Burtonsville, Maryland, USA"),
    ])

    add_paragraph(doc, " ")

    add_field_table(doc, [
        ("Receiving Party (Filing Agent)", ""),
        ("Signature", ""),
        ("Date", ""),
        ("Place", ""),
        ("Witness Name", ""),
        ("Witness Signature", ""),
    ])

    doc.save(output_path)
    print(f"  OK  {output_path.relative_to(SECURED_DIR)}")


# ---------------------------------------------------------------------------
# Step 3 — Minimal source deposit (first 25 / last 25 pages)
# ---------------------------------------------------------------------------

SOURCE_EXTENSIONS = {
    '.ts', '.tsx', '.js', '.jsx', '.mjs',
    '.prisma', '.css', '.scss', '.html', '.sql',
    '.sh', '.ps1', '.md', '.json', '.yaml', '.yml', '.example',
}

EXCLUDE_PATHS = {
    'node_modules/', '.next/', 'dist/', 'build/',
    '.vercel/', '.git/', '.turbo/', '.DS_Store',
}

EXCLUDE_FILES = {
    'package-lock.json', 'pnpm-lock.yaml', 'yarn.lock',
    'apps/web/src/locales/am.json',
    'apps/web/src/locales/om.json',
    'apps/web/src/locales/ti.json',
    'apps/web/src/locales/so.json',
}


def _should_include(file_path):
    path_str = file_path.replace('\\', '/')
    for excl in EXCLUDE_PATHS:
        if path_str.startswith(excl) or f'/{excl}' in path_str:
            return False
    basename = os.path.basename(path_str)
    if path_str in EXCLUDE_FILES or basename in EXCLUDE_FILES:
        return False
    ext = os.path.splitext(basename)[1].lower()
    return ext in SOURCE_EXTENSIONS


def _classify_volume(path):
    p = path.replace('\\', '/')
    if p.startswith('apps/web/src/app/'):
        return 1
    if p.startswith('apps/web/src/components/'):
        return 2
    if (p.startswith('apps/web/src/lib/') or
        p.startswith('apps/web/src/contexts/') or
        p.startswith('apps/web/src/hooks/') or
        p.startswith('apps/web/src/config/') or
        p.startswith('apps/web/src/locales/') or
        p.startswith('apps/web/src/styles/') or
        p.startswith('apps/web/src/types/')):
        return 3
    if p.startswith('packages/database/'):
        return 4
    return 5


def _get_ordered_source_files():
    result = subprocess.run(
        ['git', 'ls-files'],
        cwd=REPO_ROOT,
        capture_output=True,
        text=True,
        check=True,
    )
    tracked = result.stdout.strip().split('\n')
    filtered = [f for f in tracked if _should_include(f)]
    # Sort by (volume, path) — same canonical order as the full deposit.
    filtered.sort(key=lambda p: (_classify_volume(p), p))
    return filtered


def _read_file(relative):
    full = REPO_ROOT / relative
    try:
        return _sanitize_for_xml(
            full.read_text(encoding='utf-8', errors='replace')
        )
    except Exception as e:
        return f"[Could not read file: {e}]"


def _accumulate_until(files, target_lines, from_end=False):
    """Walk file list accumulating file contents until target lines reached."""
    iterator = reversed(files) if from_end else iter(files)
    bucket = []
    total = 0
    for f in iterator:
        content = _read_file(f)
        lines = content.count('\n') + (1 if content and not content.endswith('\n') else 0)
        bucket.append((f, content, lines))
        total += lines
        if total >= target_lines:
            break
    if from_end:
        bucket.reverse()
    return bucket, total


def generate_minimal_source(first_path, last_path):
    files = _get_ordered_source_files()
    print(f"  Source files considered: {len(files)}")

    first_files, first_lines = _accumulate_until(files, TARGET_LINES_PER_SIDE)
    last_files, last_lines = _accumulate_until(
        files, TARGET_LINES_PER_SIDE, from_end=True,
    )

    # Render FIRST
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "Minimal Source — First 25 Pages")
    add_title(
        doc,
        "MINIMAL SOURCE DEPOSIT — FIRST 25 PAGES",
        "Onekof Platform — Reduced Deposit per EIPA first/last rule",
    )
    add_paragraph(
        doc,
        f"This document contains the first ~25 pages of representative source "
        f"code from the Onekof Platform, rendered in the canonical volume "
        f"order (Volume 01 \u2192 Volume 05). Total: {len(first_files)} files, "
        f"{first_lines:,} lines.",
        italic=True,
    )
    for f, content, lines in first_files:
        add_file_header(doc, f, lines)
        add_code_block(doc, content)
    doc.save(first_path)
    print(f"  OK  {first_path.relative_to(SECURED_DIR)} "
          f"({len(first_files)} files, {first_lines:,} lines)")

    # Render LAST
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "Minimal Source — Last 25 Pages")
    add_title(
        doc,
        "MINIMAL SOURCE DEPOSIT — LAST 25 PAGES",
        "Onekof Platform — Reduced Deposit per EIPA first/last rule",
    )
    add_paragraph(
        doc,
        f"This document contains the last ~25 pages of representative source "
        f"code from the Onekof Platform, rendered from the end of the "
        f"canonical volume order. Total: {len(last_files)} files, "
        f"{last_lines:,} lines.",
        italic=True,
    )
    for f, content, lines in last_files:
        add_file_header(doc, f, lines)
        add_code_block(doc, content)
    doc.save(last_path)
    print(f"  OK  {last_path.relative_to(SECURED_DIR)} "
          f"({len(last_files)} files, {last_lines:,} lines)")

    # README for minimal folder
    readme = MINIMAL_DIR / "Minimal_Source_README.txt"
    readme.write_text(
        f"ONEKOF PLATFORM — MINIMAL SOURCE DEPOSIT\n"
        f"{'=' * 50}\n\n"
        f"Author: Oli Tamrat Oli <oli.oli@udc.edu>\n"
        f"Generated: {TODAY}\n\n"
        f"This folder contains the 'first 25 / last 25 pages' interpretation\n"
        f"of the source code deposit, as originally defined in EIPA Form C-1\n"
        f"Section 5 for software works exceeding 50 pages.\n\n"
        f"CONTENTS:\n"
        f"  First_25_Pages.docx  -- {len(first_files)} files, {first_lines:,} lines\n"
        f"  Last_25_Pages.docx   -- {len(last_files)} files, {last_lines:,} lines\n\n"
        f"PURPOSE:\n"
        f"  Offered as an alternative to the complete 5-volume source\n"
        f"  deposit, in case EIPA prefers the original printed-page rule\n"
        f"  over the full electronic archive.\n",
        encoding='utf-8',
    )
    print(f"  OK  {readme.relative_to(SECURED_DIR)}")


# ---------------------------------------------------------------------------
# Step 4 — Password-protect full source ZIP
# ---------------------------------------------------------------------------

def _generate_password(length=24):
    alphabet = string.ascii_letters + string.digits
    return ''.join(secrets.choice(alphabet) for _ in range(length))


def generate_encrypted_zip(password_out_path, encrypted_zip_path):
    src_zip = DEPOSIT_DIR / "onekof-platform-source.zip"
    if not src_zip.exists():
        raise RuntimeError(
            f"Expected source ZIP at {src_zip}. "
            f"Run generate-eipa-deposit.py first."
        )

    password = _generate_password()

    with pyzipper.AESZipFile(
        encrypted_zip_path,
        'w',
        compression=pyzipper.ZIP_DEFLATED,
        encryption=pyzipper.WZ_AES,
    ) as zf:
        zf.setpassword(password.encode('utf-8'))
        zf.setencryption(pyzipper.WZ_AES, nbits=256)
        zf.write(src_zip, arcname='onekof-platform-source.zip')

    password_out_path.write_text(
        f"ONEKOF PLATFORM — SECURED DEPOSIT ZIP PASSWORD\n"
        f"{'=' * 50}\n\n"
        f"DO NOT PLACE THIS FILE ON THE CD-ROM.\n"
        f"Hand this password to EIPA SEPARATELY, on paper, at the "
        f"filing desk.\n\n"
        f"File:      onekof-platform-source.zip.aes\n"
        f"Encryption: AES-256 (WinZip-compatible)\n"
        f"Password:   {password}\n\n"
        f"Generated: {TODAY}\n"
        f"Author:    Oli Tamrat Oli <oli.oli@udc.edu>\n\n"
        f"OPENING INSTRUCTIONS:\n"
        f"  Windows: Use 7-Zip or WinRAR. Double-click the .aes file\n"
        f"           and enter the password when prompted.\n"
        f"  macOS:   Use Keka or The Unarchiver.\n"
        f"  Linux:   Use 7z: 7z x onekof-platform-source.zip.aes\n",
        encoding='utf-8',
    )

    size_mb = encrypted_zip_path.stat().st_size / (1024 * 1024)
    print(f"  OK  {encrypted_zip_path.relative_to(SECURED_DIR)} ({size_mb:.1f} MB)")
    print(f"  OK  {password_out_path.relative_to(SECURED_DIR)} (KEEP OFFLINE!)")
    return password


# ---------------------------------------------------------------------------
# Step 5 — SHA-256 manifest
# ---------------------------------------------------------------------------

def _sha256_of(path):
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b''):
            h.update(chunk)
    return h.hexdigest()


def generate_sha256_manifest(txt_path, docx_path):
    # Collect every file in SECURED_DIR except the manifest files themselves
    # and the password file (which must not be hashed for the CD).
    skip_names = {
        'SHA256_MANIFEST.txt',
        'SHA256_MANIFEST.docx',
        'ZIP_PASSWORD.txt',
    }
    entries = []
    for p in sorted(SECURED_DIR.rglob('*')):
        if not p.is_file():
            continue
        if p.name in skip_names:
            continue
        rel = p.relative_to(SECURED_DIR).as_posix()
        size = p.stat().st_size
        sha = _sha256_of(p)
        entries.append((rel, size, sha))

    # Plain text
    lines = [
        "ONEKOF PLATFORM — SECURED DEPOSIT SHA-256 MANIFEST",
        "=" * 60,
        "",
        f"Generated: {TODAY}",
        f"Author:    Oli Tamrat Oli <oli.oli@udc.edu>",
        f"Files:     {len(entries)}",
        "",
        "Use this manifest to verify that no file on the CD-ROM has been",
        "modified after the deposit was generated. To verify a file:",
        "",
        "  Windows (PowerShell):",
        "    Get-FileHash -Algorithm SHA256 <filename>",
        "",
        "  Linux / macOS:",
        "    sha256sum <filename>      (Linux)",
        "    shasum -a 256 <filename>  (macOS)",
        "",
        "-" * 60,
        "",
    ]
    for rel, size, sha in entries:
        lines.append(f"{sha}  {size:>12}  {rel}")
    lines.append("")
    lines.append("-" * 60)
    lines.append("End of manifest.")
    txt_path.write_text('\n'.join(lines), encoding='utf-8')

    # Word version
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "SHA-256 Manifest")
    add_title(
        doc,
        "SHA-256 MANIFEST",
        "Onekof Platform Secured Deposit — Tamper Detection",
    )

    add_field_table(doc, [
        ("Author", "Oli Tamrat Oli"),
        ("Email", "oli.oli@udc.edu"),
        ("Generated", TODAY),
        ("Files Covered", str(len(entries))),
        ("Hash Algorithm", "SHA-256"),
    ])

    add_section_heading(doc, "Purpose")
    add_paragraph(
        doc,
        "This manifest provides a SHA-256 cryptographic hash for every file "
        "in the secured deposit. It allows the Ethiopian Intellectual "
        "Property Authority, the applicant, and any legal proceeding to "
        "verify that the contents of the deposit CD-ROM have not been "
        "altered since the date of generation. Any modification to any "
        "file — even a single byte — will produce a different hash.",
    )

    add_section_heading(doc, "Verification Commands")
    add_paragraph(doc, "Windows PowerShell:", bold=True)
    add_paragraph(doc, "    Get-FileHash -Algorithm SHA256 <filename>")
    add_paragraph(doc, "Linux:", bold=True)
    add_paragraph(doc, "    sha256sum <filename>")
    add_paragraph(doc, "macOS:", bold=True)
    add_paragraph(doc, "    shasum -a 256 <filename>")

    add_section_heading(doc, "Hash Table")

    table = doc.add_table(rows=len(entries) + 1, cols=2)
    table.style = 'Light List Accent 1'

    # Header
    hp = table.rows[0].cells[0].paragraphs[0]
    hrun = hp.add_run("File")
    hrun.bold = True
    hrun.font.size = Pt(10)
    hp2 = table.rows[0].cells[1].paragraphs[0]
    hrun2 = hp2.add_run("SHA-256 (size in bytes)")
    hrun2.bold = True
    hrun2.font.size = Pt(10)
    _shade_cell(table.rows[0].cells[0], 'E8F5F3')
    _shade_cell(table.rows[0].cells[1], 'E8F5F3')

    for i, (rel, size, sha) in enumerate(entries, start=1):
        name_cell = table.rows[i].cells[0]
        hash_cell = table.rows[i].cells[1]

        np_ = name_cell.paragraphs[0]
        nrun = np_.add_run(rel)
        nrun.font.size = Pt(8)
        nrun.font.name = 'Consolas'

        hp_ = hash_cell.paragraphs[0]
        hrun_ = hp_.add_run(f"{sha}\n{size:,} bytes")
        hrun_.font.size = Pt(7)
        hrun_.font.name = 'Consolas'

    doc.save(docx_path)
    print(f"  OK  {txt_path.relative_to(SECURED_DIR)}")
    print(f"  OK  {docx_path.relative_to(SECURED_DIR)}")
    return len(entries)


# ---------------------------------------------------------------------------
# Step 6 — README
# ---------------------------------------------------------------------------

def generate_readme(path, num_files_hashed):
    text = f"""ONEKOF PLATFORM — EIPA SECURED DEPOSIT
{'=' * 50}

Author:   Oli Tamrat Oli <oli.oli@udc.edu>
Generated: {TODAY}

This folder is an ALTERNATIVE, HARDENED version of the standard
DEPOSIT/ folder. It applies additional security measures on top of
the same underlying source code and filing forms. Offer it to EIPA
during the filing appointment and let them choose which version to
accept based on their preference.

CONTENTS:
{'-' * 50}

00_FORMS/
    Form_C-1_Application.docx          -- EIPA main application
    Software_Description.docx          -- Functional description
    Publication_Date.docx              -- First publication declaration
    Source_Code_Deposit_Cover.docx     -- CD-ROM contents cover sheet
    Filing_Agent_NDA.docx              -- NDA to be signed by the
                                          filing agent BEFORE the disc
                                          leaves the author's hands.

01_MINIMAL_SOURCE/
    First_25_Pages.docx                -- First ~25 pages of source
    Last_25_Pages.docx                 -- Last ~25 pages of source
    Minimal_Source_README.txt          -- Explanation of reduced deposit
    (These two files implement the original EIPA Form C-1 Section 5
     "first 25 / last 25 pages" rule, in case EIPA prefers the
     reduced deposit over the full 5-volume archive.)

onekof-platform-source.zip.aes         -- Full source archive,
                                          AES-256 encrypted.
                                          Password is NOT on the CD.

SHA256_MANIFEST.txt                    -- Tamper-detection hash list
SHA256_MANIFEST.docx                   -- Formatted hash list (Word)
                                          Covers {num_files_hashed} files.

README.txt                             -- This file.

NOT ON THE CD (hand separately to EIPA):
{'-' * 50}

ZIP_PASSWORD.txt                       -- The AES-256 password for the
                                          encrypted ZIP. This file MUST
                                          be kept offline, printed on
                                          paper, and handed to EIPA at
                                          the filing desk in a sealed
                                          envelope. It is saved in the
                                          DEPOSIT_SECURED/ folder on
                                          the author's local machine
                                          but SHALL NOT be burned to
                                          the CD.

USAGE:
{'-' * 50}

1. Before the filing:
   a. Have the filing agent sign Filing_Agent_NDA.docx.
   b. Print ZIP_PASSWORD.txt on paper; keep the digital copy offline.

2. Burning the CD:
   a. Burn the DEPOSIT_SECURED/ folder to a CD-R (write-once media).
   b. DO NOT include ZIP_PASSWORD.txt on the CD.
   c. Verify the burned disc with SHA256_MANIFEST.txt before handing
      it to the filing agent.

3. At the EIPA filing desk:
   a. Present the standard DEPOSIT/ first. If EIPA accepts it, use it.
   b. If EIPA has concerns about the completeness, size, or sensitivity
      of the standard deposit, offer this DEPOSIT_SECURED/ alternative.
   c. If EIPA prefers the reduced deposit, point them to
      01_MINIMAL_SOURCE/.
   d. If EIPA accepts the encrypted archive, hand them the paper
      ZIP_PASSWORD.txt separately.

4. After the filing:
   a. Obtain the EIPA-stamped receipt.
   b. Rotate all production secrets (NEXTAUTH_SECRET, database
      password, API keys) as a precaution.
   c. Retain the signed NDA copy and the filing receipt together.

SECURITY PROPERTIES:
{'-' * 50}

  (a) AES-256 encrypted ZIP        -- Confidentiality of full archive
  (b) SHA-256 manifest             -- Tamper detection
  (c) Minimal first/last deposit   -- Reduces attack surface if
                                      the full deposit is not required
  (d) Filing Agent NDA             -- Legal protection against
                                      unauthorized copying
  (e) Separate password channel    -- Defense in depth: even if the CD
                                      is stolen, the encrypted ZIP
                                      cannot be opened without the
                                      separately-delivered password
"""
    path.write_text(text, encoding='utf-8')
    print(f"  OK  {path.relative_to(SECURED_DIR)}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("ONEKOF — EIPA SECURED DEPOSIT GENERATOR")
    print("=" * 60)

    if SECURED_DIR.exists():
        print(f"Removing existing {SECURED_DIR}...")
        shutil.rmtree(SECURED_DIR)
    SECURED_DIR.mkdir(parents=True)
    MINIMAL_DIR.mkdir(parents=True)

    print()
    print("[1/6] Copying existing EIPA forms from DEPOSIT/00_FORMS/...")
    copied = copy_existing_forms()
    print(f"      {copied} forms copied")

    print()
    print("[2/6] Generating Filing Agent NDA...")
    generate_filing_agent_nda(FORMS_DIR / "Filing_Agent_NDA.docx")

    print()
    print("[3/6] Generating minimal source deposit (first/last 25 pages)...")
    generate_minimal_source(
        MINIMAL_DIR / "First_25_Pages.docx",
        MINIMAL_DIR / "Last_25_Pages.docx",
    )

    print()
    print("[4/6] Encrypting full source ZIP with AES-256...")
    password = generate_encrypted_zip(
        SECURED_DIR / "ZIP_PASSWORD.txt",
        SECURED_DIR / "onekof-platform-source.zip.aes",
    )

    print()
    print("[5/6] Generating SHA-256 manifest...")
    num_hashed = generate_sha256_manifest(
        SECURED_DIR / "SHA256_MANIFEST.txt",
        SECURED_DIR / "SHA256_MANIFEST.docx",
    )

    print()
    print("[6/6] Writing README...")
    generate_readme(SECURED_DIR / "README.txt", num_hashed)

    print()
    print("=" * 60)
    print("SECURED DEPOSIT READY")
    print("=" * 60)
    print(f"Location: {SECURED_DIR}")
    print()
    print("IMPORTANT: ZIP_PASSWORD.txt MUST NOT be burned to the CD.")
    print("           Print it on paper and keep it offline.")
    print()


if __name__ == '__main__':
    main()
