"""
Re-encrypt the DEPOSIT_SECURED ZIP with a user-chosen password and
regenerate the SHA-256 manifest to reflect the new ZIP hash.

Run this any time the password changes.
"""

import hashlib
import sys
from datetime import datetime
from pathlib import Path

import pyzipper
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).parent
DEPOSIT_DIR = REPO_ROOT / "DEPOSIT"
SECURED_DIR = REPO_ROOT / "DEPOSIT_SECURED"
PASSWORD_FILE = Path("C:/Users/olita/EIPA_ZIP_PASSWORD_KEEP_OFFLINE.txt")

TEAL = RGBColor(0x1C, 0x8C, 0x7D)
DARK = RGBColor(0x22, 0x27, 0x2B)
GREY = RGBColor(0x77, 0x77, 0x77)

TODAY = datetime.now().strftime("%B %d, %Y")


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_document_defaults(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = TEAL
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1C8C7D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True


def add_field_table(doc, rows, label_width_inches=2.3):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.style = 'Light List Accent 1'
    for i, (label, value) in enumerate(rows):
        lc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        lc.width = Inches(label_width_inches)
        vc.width = Inches(6.5 - label_width_inches)
        _shade_cell(lc, 'E8F5F3')
        lp = lc.paragraphs[0]
        lrun = lp.add_run(label)
        lrun.bold = True
        lrun.font.size = Pt(10)
        lrun.font.color.rgb = DARK
        vp = vc.paragraphs[0]
        display = value if value else "__________________________________________"
        vrun = vp.add_run(display)
        vrun.font.size = Pt(10)
        if not value:
            vrun.font.color.rgb = GREY
    doc.add_paragraph()


def add_footer(doc, label):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{label}  |  Onekof Platform — SECURED DEPOSIT  |  "
        f"Oli Tamrat Oli  |  {TODAY}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


def sha256_of(path):
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b''):
            h.update(chunk)
    return h.hexdigest()


def reencrypt(password):
    src_zip = DEPOSIT_DIR / "onekof-platform-source.zip"
    enc_zip = SECURED_DIR / "onekof-platform-source.zip.aes"

    if enc_zip.exists():
        enc_zip.unlink()

    with pyzipper.AESZipFile(
        enc_zip, 'w',
        compression=pyzipper.ZIP_DEFLATED,
        encryption=pyzipper.WZ_AES,
    ) as zf:
        zf.setpassword(password.encode('utf-8'))
        zf.setencryption(pyzipper.WZ_AES, nbits=256)
        zf.write(src_zip, arcname='onekof-platform-source.zip')

    size_mb = enc_zip.stat().st_size / (1024 * 1024)
    print(f"  OK  onekof-platform-source.zip.aes ({size_mb:.1f} MB)")
    return enc_zip


def write_password_file(password):
    PASSWORD_FILE.write_text(
        f"ONEKOF PLATFORM — SECURED DEPOSIT ZIP PASSWORD\n"
        f"{'=' * 50}\n\n"
        f"DO NOT PLACE THIS FILE ON THE CD-ROM.\n"
        f"Print this page, seal it in an envelope, and hand it to\n"
        f"EIPA at the filing desk.\n\n"
        f"File:       onekof-platform-source.zip.aes\n"
        f"Encryption: AES-256 (WinZip-compatible)\n"
        f"Password:   {password}\n\n"
        f"Generated:  {TODAY}\n"
        f"Author:     Oli Tamrat Oli <oli.oli@udc.edu>\n\n"
        f"OPENING INSTRUCTIONS FOR EIPA CLERK:\n"
        f"-------------------------------------\n"
        f"1. Insert the CD-ROM into any Windows 10 or later PC.\n"
        f"2. Copy 'onekof-platform-source.zip.aes' to the desktop.\n"
        f"3. Rename it to 'onekof-platform-source.zip' (remove '.aes').\n"
        f"4. Right-click the file and choose 'Extract All...'.\n"
        f"5. When Windows asks for a password, enter exactly:\n\n"
        f"      {password}\n\n"
        f"   (Note the capital letters, digits, dash, and '@' symbol.)\n\n"
        f"6. If Windows's built-in extractor does not accept AES-256,\n"
        f"   use the free tool 7-Zip instead:\n"
        f"   https://www.7-zip.org\n"
        f"   Right-click the .aes file, choose 7-Zip > Extract Here,\n"
        f"   and enter the same password.\n\n"
        f"The extracted folder contains the complete, human-readable\n"
        f"source code of the Onekof Platform.\n",
        encoding='utf-8',
    )
    print(f"  OK  {PASSWORD_FILE.name} (kept OUTSIDE DEPOSIT_SECURED/)")


def regenerate_manifest():
    skip_names = {
        'SHA256_MANIFEST.txt',
        'SHA256_MANIFEST.docx',
        'ZIP_PASSWORD.txt',
    }
    entries = []
    for p in sorted(SECURED_DIR.rglob('*')):
        if not p.is_file():
            continue
        if p.name in skip_names:
            continue
        rel = p.relative_to(SECURED_DIR).as_posix()
        size = p.stat().st_size
        sha = sha256_of(p)
        entries.append((rel, size, sha))

    # plain text
    lines = [
        "ONEKOF PLATFORM — SECURED DEPOSIT SHA-256 MANIFEST",
        "=" * 60,
        "",
        f"Generated: {TODAY}",
        f"Author:    Oli Tamrat Oli <oli.oli@udc.edu>",
        f"Files:     {len(entries)}",
        "",
        "Use this manifest to verify that no file on the CD-ROM has been",
        "modified after the deposit was generated. To verify a file:",
        "",
        "  Windows (PowerShell):",
        "    Get-FileHash -Algorithm SHA256 <filename>",
        "",
        "  Linux / macOS:",
        "    sha256sum <filename>      (Linux)",
        "    shasum -a 256 <filename>  (macOS)",
        "",
        "-" * 60,
        "",
    ]
    for rel, size, sha in entries:
        lines.append(f"{sha}  {size:>12}  {rel}")
    lines.append("")
    lines.append("-" * 60)
    lines.append("End of manifest.")
    (SECURED_DIR / "SHA256_MANIFEST.txt").write_text(
        '\n'.join(lines), encoding='utf-8')

    # word
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc, "SHA-256 Manifest")
    add_title(
        doc,
        "SHA-256 MANIFEST",
        "Onekof Platform Secured Deposit — Tamper Detection",
    )
    add_field_table(doc, [
        ("Author", "Oli Tamrat Oli"),
        ("Email", "oli.oli@udc.edu"),
        ("Generated", TODAY),
        ("Files Covered", str(len(entries))),
        ("Hash Algorithm", "SHA-256"),
    ])
    add_section_heading(doc, "Purpose")
    add_paragraph(
        doc,
        "This manifest provides a SHA-256 cryptographic hash for every file "
        "in the secured deposit. It allows the Ethiopian Intellectual "
        "Property Authority, the applicant, and any legal proceeding to "
        "verify that the contents of the deposit CD-ROM have not been "
        "altered since the date of generation. Any modification to any "
        "file — even a single byte — will produce a different hash.",
    )
    add_section_heading(doc, "Verification Commands")
    add_paragraph(doc, "Windows PowerShell:", bold=True)
    add_paragraph(doc, "    Get-FileHash -Algorithm SHA256 <filename>")
    add_paragraph(doc, "Linux:", bold=True)
    add_paragraph(doc, "    sha256sum <filename>")
    add_paragraph(doc, "macOS:", bold=True)
    add_paragraph(doc, "    shasum -a 256 <filename>")
    add_section_heading(doc, "Hash Table")

    table = doc.add_table(rows=len(entries) + 1, cols=2)
    table.style = 'Light List Accent 1'
    hp = table.rows[0].cells[0].paragraphs[0]
    hrun = hp.add_run("File")
    hrun.bold = True
    hrun.font.size = Pt(10)
    hp2 = table.rows[0].cells[1].paragraphs[0]
    hrun2 = hp2.add_run("SHA-256 (size in bytes)")
    hrun2.bold = True
    hrun2.font.size = Pt(10)
    _shade_cell(table.rows[0].cells[0], 'E8F5F3')
    _shade_cell(table.rows[0].cells[1], 'E8F5F3')

    for i, (rel, size, sha) in enumerate(entries, start=1):
        nc = table.rows[i].cells[0]
        hc = table.rows[i].cells[1]
        nrun = nc.paragraphs[0].add_run(rel)
        nrun.font.size = Pt(8)
        nrun.font.name = 'Consolas'
        hrun_ = hc.paragraphs[0].add_run(f"{sha}\n{size:,} bytes")
        hrun_.font.size = Pt(7)
        hrun_.font.name = 'Consolas'

    doc.save(SECURED_DIR / "SHA256_MANIFEST.docx")
    print(f"  OK  SHA256_MANIFEST.txt")
    print(f"  OK  SHA256_MANIFEST.docx")
    return len(entries)


def main():
    if len(sys.argv) < 2:
        print("Usage: python reencrypt-with-password.py <password>")
        sys.exit(1)
    password = sys.argv[1]

    print("=" * 60)
    print("RE-ENCRYPT SECURED DEPOSIT")
    print("=" * 60)
    print()
    print("[1/3] Re-encrypting ZIP with new password...")
    enc = reencrypt(password)

    # Verify by opening
    with pyzipper.AESZipFile(enc) as zf:
        zf.setpassword(password.encode('utf-8'))
        names = zf.namelist()
        with zf.open(names[0]) as inner:
            sig = inner.read(4)
        assert sig == b'PK\x03\x04', f"Inner ZIP signature invalid: {sig}"
    print("      Verified: encrypted ZIP opens with new password.")

    print()
    print("[2/3] Writing password file (OUTSIDE DEPOSIT_SECURED/)...")
    write_password_file(password)

    print()
    print("[3/3] Regenerating SHA-256 manifest...")
    n = regenerate_manifest()
    print(f"      {n} files hashed.")

    print()
    print("=" * 60)
    print("DONE.")
    print("=" * 60)
    print(f"Password file: {PASSWORD_FILE}")
    print(f"Encrypted ZIP: {enc}")


if __name__ == '__main__':
    main()
